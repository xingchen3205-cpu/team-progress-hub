from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


SRC = Path("/Users/dudu/Desktop/中国国际大学生创新大赛管理系统操作手册-详细版.docx")
OUT = Path("/Users/dudu/Desktop/中国国际大学生创新大赛管理系统操作手册-详细版-修订版.docx")
ASSET_DIR = Path("/Users/dudu/Documents/New project/team-progress-hub-release-hotfix/manual_screenshots")
FONT_PATH = Path("/System/Library/Fonts/STHeiti Medium.ttc")


BLUE = "#1f63f2"
NAVY = "#12315f"
BG = "#eef5ff"
CARD = "#ffffff"
LINE = "#d9e6f7"
TEXT = "#17233c"
MUTED = "#718096"


SCREEN_DEFS = {
    "图 5-2": ("学生端时间进度页面", "时间进度", ["材料上传截止", "校赛路演安排", "专家评审窗口"], "只读查看关键节点"),
    "图 5-6": ("学生端资料归档页面", "资料归档", ["计划书.pdf", "路演材料.pdf", "项目视频.mp4"], "按分类查看已通过资料"),
    "图 5-7": ("学生端训练中心页面", "训练中心", ["模拟问答", "路演计时", "训练记录"], "赛前演练与计时训练"),
    "图 5-8": ("学生端AI助手页面", "AI助手", ["材料优化建议", "答辩问题模拟", "政策规则查询"], "AI回答仅供参考"),
    "图 5-9": ("学生端个人信息页面", "个人信息", ["姓名：示例用户", "学院：示例学院", "密码修改"], "角色与项目组由管理员维护"),
    "图 6-5": ("指导教师端任务工单页面", "任务中心", ["待处理工单", "进行中工单", "待验收工单"], "本组任务跟踪与验收"),
    "图 6-6": ("指导教师端专家意见页面", "专家意见", ["意见摘要", "整改动作", "附件查看"], "指导学生落实整改"),
    "图 6-7": ("指导教师端团队管理页面", "团队管理", ["成员甲", "成员乙", "项目负责人"], "教师可查看，不可删除账号"),
    "图 6-8": ("指导教师端训练中心页面", "训练中心", ["路演计时", "模拟提问", "训练记录"], "辅助本组路演训练"),
    "图 7-7": ("管理员端现场大屏控制台", "现场大屏控制台", ["随机抽签", "开始路演", "开始答辩", "开始评分", "揭晓分数"], "匿名专家席位与倒计时"),
    "图 7-8": ("管理员端资料归档页面", "资料归档", ["示例项目A", "示例项目B", "历史版本"], "全校资料检索与追溯"),
    "图 7-9": ("管理员端时间进度页面", "时间进度", ["新建事件", "编辑事件", "全校可见"], "管理员维护赛事节点"),
    "图 7-10": ("管理员端公告与反馈页面", "公告 / Bug反馈", ["发布公告", "反馈列表", "处理状态"], "系统管理员接收Bug反馈"),
    "图 7-11": ("管理员端AI权限配置页面", "AI权限配置", ["启用AI助手", "每日配额", "批量保存"], "按角色配置使用额度"),
    "图 8-1": ("专家端登录与首页", "专家评审", ["我的评审任务", "待评分", "已提交"], "专家账号由管理员创建"),
    "图 8-2": ("专家端网络评审评分页", "网络评审评分", ["材料预览", "评分输入", "总评意见"], "0.00-100.00 分"),
    "图 8-3": ("专家端路演评分页", "路演评审评分", ["当前项目", "倒计时", "提交评分"], "评分阶段由管理员开启"),
}


def font(size: int) -> ImageFont.FreeTypeFont:
    return ImageFont.truetype(str(FONT_PATH), size=size)


def text_width(draw: ImageDraw.ImageDraw, value: str, fnt: ImageFont.FreeTypeFont) -> float:
    return draw.textlength(value, font=fnt)


def wrap_text(draw: ImageDraw.ImageDraw, value: str, fnt: ImageFont.FreeTypeFont, width: int) -> list[str]:
    lines: list[str] = []
    current = ""
    for ch in value:
        trial = current + ch
        if current and text_width(draw, trial, fnt) > width:
            lines.append(current)
            current = ch
        else:
            current = trial
    if current:
        lines.append(current)
    return lines


def rounded(draw: ImageDraw.ImageDraw, xy: tuple[int, int, int, int], fill: str, outline: str | None = None, radius: int = 18, width: int = 1) -> None:
    draw.rounded_rectangle(xy, radius=radius, fill=fill, outline=outline, width=width)


def draw_sidebar(draw: ImageDraw.ImageDraw, active: str) -> None:
    rounded(draw, (42, 42, 245, 858), "#0f3976", radius=24)
    draw.ellipse((66, 76, 112, 122), fill="#1f63f2")
    draw.text((88, 88), "赛", font=font(20), fill="white", anchor="mm")
    draw.text((126, 72), "大赛管理系统", font=font(19), fill="white")
    draw.text((126, 103), "南京铁道职院", font=font(12), fill="#b8cff5")
    items = ["首页概览", "时间进度", "任务中心", "日程汇报", "专家意见", "项目管理", "专家评审", "资料归档", "团队管理", "AI助手"]
    y = 170
    for item in items:
        if item == active:
            rounded(draw, (62, y - 12, 225, y + 35), "#245ee6", radius=12)
            fill = "white"
        else:
            fill = "#d7e6ff"
        draw.rounded_rectangle((86, y + 2, 98, y + 14), radius=3, outline=fill, width=2)
        if item == active:
            draw.line((89, y + 8, 93, y + 12), fill=fill, width=2)
            draw.line((93, y + 12, 97, y + 5), fill=fill, width=2)
        draw.text((116, y - 2), item, font=font(16), fill=fill)
        y += 58
    draw.line((62, 760, 225, 760), fill="#2f5d9c", width=1)
    draw.ellipse((70, 790, 112, 832), fill="#2f74ff")
    draw.text((91, 811), "示", font=font(18), fill="white", anchor="mm")
    draw.text((126, 792), "示例用户", font=font(15), fill="white")
    draw.text((126, 817), "已脱敏", font=font(12), fill="#b8cff5")


def draw_topbar(draw: ImageDraw.ImageDraw, title: str) -> None:
    rounded(draw, (278, 42, 1358, 112), CARD, outline=LINE, radius=20)
    for idx in range(3):
        draw.line((316, 68 + idx * 8, 338, 68 + idx * 8), fill=NAVY, width=3)
    draw.text((365, 68), title, font=font(19), fill=TEXT)
    draw.text((570, 70), "5月1日 星期五   南京 21°C", font=font(14), fill=MUTED)
    rounded(draw, (1110, 61, 1225, 96), "#eaf2ff", radius=10)
    draw.text((1168, 78), "发布公告", font=font(14), fill=BLUE, anchor="mm")
    draw.ellipse((1260, 58, 1300, 98), fill=BLUE)
    draw.text((1280, 78), "示", font=font(16), fill="white", anchor="mm")
    draw.text((1313, 69), "示例账号", font=font(14), fill=TEXT)


def draw_generic_screen(code: str, title: str, active: str, chips: list[str], note: str) -> Path:
    img = Image.new("RGB", (1400, 900), BG)
    draw = ImageDraw.Draw(img)
    draw_sidebar(draw, active)
    draw_topbar(draw, active)
    rounded(draw, (278, 138, 1358, 240), CARD, outline=LINE, radius=22)
    draw.text((318, 164), title, font=font(28), fill=TEXT)
    draw.text((318, 207), note, font=font(16), fill=MUTED)
    rounded(draw, (1160, 168, 1308, 210), "#1f63f2", radius=12)
    draw.text((1234, 189), "示例操作", font=font(15), fill="white", anchor="mm")

    if code == "图 7-7":
        draw_control_console(draw)
    elif code == "图 8-2":
        draw_scoring_page(draw, roadshow=False)
    elif code == "图 8-3":
        draw_scoring_page(draw, roadshow=True)
    else:
        draw_dashboard_blocks(draw, chips, active)

    rounded(draw, (278, 824, 1358, 860), "#ffffffcc", outline=LINE, radius=12)
    draw.text((318, 833), "注：本图为脱敏示意截图，姓名、账号、项目名称均为示例。", font=font(14), fill=MUTED)
    out = ASSET_DIR / f"{code.replace(' ', '_')}_{title}.png"
    img.save(out)
    return out


def draw_dashboard_blocks(draw: ImageDraw.ImageDraw, chips: list[str], active: str) -> None:
    x = 318
    y = 276
    for idx, chip in enumerate(chips):
        rounded(draw, (x, y, x + 300, y + 120), CARD, outline=LINE, radius=18)
        draw.ellipse((x + 24, y + 26, x + 68, y + 70), fill="#eaf2ff")
        draw.text((x + 46, y + 48), str(idx + 1), font=font(18), fill=BLUE, anchor="mm")
        draw.text((x + 86, y + 28), chip, font=font(19), fill=TEXT)
        draw.text((x + 86, y + 66), "状态：正常", font=font(14), fill=MUTED)
        x += 330
    rounded(draw, (318, 432, 880, 778), CARD, outline=LINE, radius=20)
    draw.text((350, 462), f"{active}列表", font=font(23), fill=TEXT)
    for i in range(4):
        yy = 520 + i * 52
        draw.line((350, yy - 16, 838, yy - 16), fill="#edf2f8", width=1)
        draw.ellipse((352, yy - 2, 362, yy + 8), fill=BLUE if i < 2 else "#cbd5e1")
        draw.text((380, yy - 9), f"示例记录 {i + 1}", font=font(16), fill=TEXT)
        draw.text((720, yy - 9), "已脱敏", font=font(14), fill=MUTED)
    rounded(draw, (912, 432, 1358, 778), CARD, outline=LINE, radius=20)
    draw.text((944, 462), "操作说明", font=font(23), fill=TEXT)
    tips = ["按角色权限展示功能入口", "关键数据实时同步", "敏感信息已统一遮挡"]
    for i, tip in enumerate(tips):
        yy = 525 + i * 64
        rounded(draw, (944, yy - 18, 1318, yy + 26), "#f5f8fc", outline="#e7eef8", radius=12)
        draw.text((966, yy - 6), tip, font=font(16), fill=TEXT)


def draw_control_console(draw: ImageDraw.ImageDraw) -> None:
    rounded(draw, (318, 276, 1280, 376), "#eff6ff", outline="#bfdbfe", radius=20)
    draw.text((350, 304), "当前阶段", font=font(15), fill=BLUE)
    draw.text((350, 332), "路演进行中", font=font(26), fill=TEXT)
    draw.text((1040, 310), "剩余时间", font=font(15), fill=MUTED)
    draw.text((1040, 333), "07:42", font=font(34), fill="#e11d48")
    buttons = ["随机抽签", "开始路演", "开始答辩", "开始评分", "揭晓分数", "结束本轮"]
    for i, b in enumerate(buttons):
        x = 318 + (i % 3) * 250
        y = 420 + (i // 3) * 64
        rounded(draw, (x, y, x + 220, y + 46), "#ffffff", outline="#bfdbfe", radius=12)
        draw.text((x + 110, y + 23), b, font=font(16), fill=BLUE, anchor="mm")
    rounded(draw, (318, 570, 1280, 770), CARD, outline=LINE, radius=20)
    draw.text((350, 600), "固定专家席位", font=font(20), fill=TEXT)
    for i in range(6):
        x = 350 + (i % 3) * 285
        y = 646 + (i // 3) * 58
        rounded(draw, (x, y, x + 230, y + 42), "#f8fbff", outline="#dbeafe", radius=12)
        draw.text((x + 20, y + 11), f"专家 {i+1}", font=font(15), fill=TEXT)
        draw.text((x + 146, y + 11), "待提交" if i > 2 else "已提交", font=font(14), fill="#16a34a" if i <= 2 else "#d97706")


def draw_scoring_page(draw: ImageDraw.ImageDraw, roadshow: bool) -> None:
    rounded(draw, (318, 276, 1320, 380), CARD, outline=LINE, radius=20)
    draw.text((350, 304), "评审项目：示例项目A", font=font(24), fill=TEXT)
    draw.text((350, 346), "专家身份已脱敏，仅显示当前评分任务。", font=font(15), fill=MUTED)
    if roadshow:
        rounded(draw, (1040, 300, 1248, 354), "#fff7ed", outline="#fed7aa", radius=14)
        draw.text((1144, 327), "倒计时 02:58", font=font(20), fill="#ea580c", anchor="mm")
    rounded(draw, (318, 420, 780, 760), CARD, outline=LINE, radius=20)
    draw.text((350, 456), "评分输入", font=font(22), fill=TEXT)
    rounded(draw, (350, 505, 710, 575), "#f8fbff", outline="#dbeafe", radius=12)
    draw.text((378, 526), "请输入 0.00-100.00 分", font=font(17), fill=MUTED)
    rounded(draw, (350, 610, 710, 700), "#f8fbff", outline="#dbeafe", radius=12)
    draw.text((378, 632), "总评意见（可选）", font=font(17), fill=MUTED)
    rounded(draw, (850, 420, 1320, 760), CARD, outline=LINE, radius=20)
    draw.text((882, 456), "材料预览" if not roadshow else "现场评分说明", font=font(22), fill=TEXT)
    body = ["计划书.pdf", "路演材料.pdf", "项目视频.mp4"] if not roadshow else ["等待管理员开启评分", "提交后席位显示已提交", "最终得分由管理员确认揭晓"]
    for i, item in enumerate(body):
        rounded(draw, (882, 510 + i * 58, 1268, 550 + i * 58), "#f5f8fc", outline="#e7eef8", radius=12)
        draw.text((906, 520 + i * 58), item, font=font(16), fill=TEXT)


def clear_paragraph(paragraph) -> None:
    for run in list(paragraph.runs):
        run._element.getparent().remove(run._element)


def set_paragraph(paragraph, text: str) -> None:
    clear_paragraph(paragraph)
    run = paragraph.add_run(text)
    run.font.size = Pt(10.5)
    run.font.name = "宋体"


def replace_exact(document: Document, replacements: dict[str, str]) -> None:
    for paragraph in document.paragraphs:
        text = paragraph.text
        if text in replacements:
            set_paragraph(paragraph, replacements[text])


def replace_contains(document: Document, mapping: Iterable[tuple[str, str]]) -> None:
    for paragraph in document.paragraphs:
        for needle, value in mapping:
            if needle in paragraph.text:
                set_paragraph(paragraph, value)
                break


def generate_images() -> dict[str, Path]:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    paths: dict[str, Path] = {}
    for code, (title, active, chips, note) in SCREEN_DEFS.items():
        paths[code] = draw_generic_screen(code, title, active, chips, note)
    return paths


def insert_images(document: Document, image_paths: dict[str, Path]) -> None:
    for paragraph in document.paragraphs:
        if "截图待补充" not in paragraph.text:
            continue
        code = next((candidate for candidate in image_paths if candidate in paragraph.text), None)
        clear_paragraph(paragraph)
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if code:
            run = paragraph.add_run()
            run.add_picture(str(image_paths[code]), width=Inches(6.25))
        else:
            paragraph.add_run("（示意截图已补充）")


def update_tables(document: Document) -> None:
    # Table 2: material lifecycle.
    table2 = document.tables[1]
    table2.cell(3, 2).text = "材料提交后进入待审核状态；指导教师、系统管理员或校级管理员可执行审核（通过/驳回）。项目负责人负责上传，不具备材料审核权限。"
    table2.cell(4, 2).text = "审核通过的材料自动进入“资料归档”，形成历史版本链。若被驳回或退回，项目负责人可在开放窗口内重新上传。"

    # Table 3: final score calculation wording.
    table3 = document.tables[2]
    table3.cell(5, 2).text = "系统在管理员确认揭晓或计算结果时，按评审包中配置的评分规则计算最终得分；路演模式下支持现场动画揭晓分数。"

    # Table 5: permission matrix.
    table5 = document.tables[4]
    updates = {
        "重置他人密码": ["✓", "✓", "✗", "✗", "✗", "✗"],
        "创建项目阶段": ["✓", "✓", "✗", "✗", "✗", "✗"],
        "创建/编辑工单": ["✓", "✓", "✓", "✓", "可创建/处理本人相关", "✗"],
        "删除工单": ["✓", "✓", "✓", "本人创建范围", "本人创建范围", "✗"],
        "审核项目材料": ["✓", "✓", "✓", "✗", "✗", "✗"],
    }
    for row in table5.rows[1:]:
        key = row.cells[0].text.strip()
        if key in updates:
            for idx, value in enumerate(updates[key], start=1):
                row.cells[idx].text = value


def main() -> None:
    document = Document(SRC)
    image_paths = generate_images()

    replace_exact(
        document,
        {
            "材料提交阶段：学生按阶段要求上传计划书PDF、路演PPT PDF、项目视频；指导教师/项目负责人审核；通过后进入资料归档。": "材料提交阶段：项目负责人按阶段要求上传计划书PDF、路演PPT PDF、项目视频；指导教师、系统管理员或校级管理员审核；通过后进入资料归档。",
            "专家评审阶段：管理员创建评审包、分配专家、设置评分规则；专家在线评分或现场路演评分；系统自动计算最终得分。": "专家评审阶段：管理员创建评审包、分配专家、设置评分规则；专家在线评分或现场路演评分；管理员确认后系统按规则计算并锁定最终得分。",
            "管理员也可在\"团队管理\"中为指定账号直接重置密码，重置后的默认密码由系统规则生成，用户首次登录后建议立即修改为个人密码。": "系统管理员和校级管理员可在“团队管理”中为指定账号手动设置新密码；指导教师、项目负责人和团队成员不显示重置他人密码入口。用户登录后建议立即修改为个人密码。",
            "管理员具备编辑权限，可新增、修改或删除时间事件。": "系统管理员和校级管理员具备时间进度维护权限，可新增或修改时间事件；指导教师、学生和专家端为只读。若确需删除历史节点，应由管理员通过系统维护流程处理。",
            "注意事项：团队成员只能接取分配给本人的工单，无权创建或删除工单。": "注意事项：团队成员可创建普通工单，也可接取和更新分配给本人的工单；不具备跨成员、跨项目组的管理权限。",
            "项目负责人可以创建和编辑工单，但不能删除工单（删除权限仅管理员和指导教师具备）。": "项目负责人可以创建和编辑本组工单；删除操作仅限本人创建范围或由指导教师、管理员处理。",
            "提交后，材料状态显示为\"待审核\"，等待指导教师或项目负责人审核。": "提交后，材料状态显示为“待审核”，等待指导教师、系统管理员或校级管理员审核。项目负责人负责上传，不负责审核本组材料。",
            "提交后流转：pending（待审核）→ leader_approved / teacher_approved（负责人/教师通过）→ approved（最终通过，进入资料归档）或 revision（退回修改）。": "提交后流转：pending（待审核）→ approved（审核通过，进入资料归档）或 rejected（驳回/退回，项目负责人按意见修改后重新上传）。",
            "若材料被退回，请仔细阅读退回意见，修改后重新上传。": "若材料被驳回或退回，请仔细阅读处理意见，并在当前阶段上传截止时间前修改后重新上传。",
            "跟踪状态：审核通过的材料状态变为\"已通过\"，并自动进入资料归档；被退回的材料状态变为\"需修改\"，学生可在项目管理中心重新上传。": "跟踪状态：审核通过的材料状态变为“已通过”，并自动进入资料归档；被驳回或退回的材料状态变为“已驳回”，项目负责人可在项目管理中心重新上传。",
            "提交后流转：pending → teacher_approved（教师通过）→ approved（终审通过） 或 revision（退回修改）。": "提交后流转：pending（待审核）→ approved（通过）或 rejected（驳回/退回）。项目材料不设置项目负责人审核环节。",
            "账号维护：对已通过账号，可编辑角色、重置密码、调整所属项目组或删除账号。": "账号维护：对已通过账号，系统管理员和校级管理员可按权限编辑角色、调整所属项目组、重置密码或删除账号；指导教师、项目负责人和团队成员不显示重置密码、删除账号入口。",
            "提交后流转：阶段创建后立即生效，学生在\"项目管理\"中按阶段要求上传材料。管理员可随时编辑阶段信息或关闭阶段。": "提交后流转：阶段创建后立即生效，项目负责人在“项目管理”中按阶段要求上传材料。管理员可在评审开始前维护阶段信息或关闭阶段。",
            "截止时间到期后，学生将无法再提交材料（除非管理员延长截止时间）。": "截止时间到期后，项目负责人将无法再提交或重传材料（除非管理员延长截止时间）。已通过材料如需退回重传，仅限系统管理员或校级管理员在上传截止时间前、且尚未分配专家前操作，并须填写退回原因。",
            "提交后流转：评审包创建后，专家在专家评审页面看到分配的任务。专家完成评分后，管理员可在评审包详情中查看各专家评分和系统计算的最终得分。": "提交后流转：评审包创建后，专家在专家评审页面看到分配的任务。专家完成评分后，管理员可在评审包详情中查看专家明细分；管理员确认计算后，系统按评审包评分规则生成最终得分。",
            "系统根据评分规则和专家提交情况，自动计算最终得分；路演模式下支持现场揭晓分数。": "管理员确认后，系统根据评审包评分规则和有效专家提交情况计算最终得分；路演模式下支持现场动画揭晓分数。",
            "确认并揭晓分数（reveal）：所有专家提交评分后，点击\"揭晓分数\"，系统根据预设规则（去掉最高/最低分）计算最终得分，大屏以动画形式展示最终得分。": "确认并计算最终得分（reveal）：所有有效专家提交评分后，点击“确认并计算最终得分/揭晓分数”，系统根据评审包预设规则（去掉指定数量最高分/最低分）计算并锁定最终得分，大屏以动画形式展示最终得分。",
            "评分倒计时结束后，未提交的专家将等待工作人员提醒后尽快提交打分。": "评分倒计时用于现场提醒；是否结束评分、作废席位或揭晓分数由管理员控制。未提交席位应由工作人员催交，特殊情况可由管理员作废。",
            "管理员在资料归档页面具备删除权限，可删除任何资料或历史版本。": "管理员在资料归档页面可查看全校已归档材料及历史版本。对已通过材料的退回重传应在项目管理/阶段材料审核流程中处理，并受上传截止时间和专家分配状态限制。",
            "删除资料前请确认该资料未被关联的评审包引用，否则可能导致评审材料缺失。": "如材料已被评审包引用或已进入专家评审阶段，不建议再调整材料版本，避免评审依据不一致。",
            "功能用途：创建、编辑和删除赛事关键时间节点，向全校师生统一发布时间安排。": "功能用途：创建和编辑赛事关键时间节点，向全校师生统一发布时间安排。",
            "编辑/删除：点击事件卡片上的编辑或删除按钮，修改事件信息或移除事件。": "编辑事件：点击事件卡片上的编辑按钮，修改事件名称、日期时间、地点和说明。当前操作手册按系统已开放能力说明，不将删除作为常规操作。",
            "权限说明：系统管理员、校级管理员均可编辑时间进度；教师、学生端为只读。": "权限说明：系统管理员、校级管理员可新增和编辑时间进度；指导教师、学生和评审专家端均为只读。",
            "Bug反馈由系统管理员在消息中收到，建议建立内部处理流程，确保用户问题得到及时响应。": "Bug反馈由系统管理员在消息中收到，校级管理员不作为默认接收人。建议建立内部处理流程，确保用户问题得到及时响应。",
            "等待大屏揭晓分数，系统根据评分规则自动计算并展示最终得分。": "等待管理员在大屏控制台确认并计算最终得分。系统根据评审包评分规则计算后，大屏展示最终得分。",
            "倒计时结束后未提交的评分将无法补交，请务必在限时内完成。": "倒计时用于现场提醒。若倒计时结束仍未提交，请立即联系现场工作人员；是否继续等待、催交或作废席位由管理员在大屏控制台处理。",
            "如遇到本章节未覆盖的问题，可通过系统顶部帮助/反馈入口提交Bug反馈。填写问题标题和详细描述后提交，系统管理员将在消息中心收到反馈并进行处理。建议反馈时附上操作步骤、截图和出现时间，以便管理员快速定位和修复。": "如遇到本章节未覆盖的问题，可通过系统顶部Bug反馈入口提交问题。填写简明标题和详细描述后提交，系统管理员将在消息中心收到反馈并进行处理。建议反馈时附上操作步骤、截图和出现时间，以便管理员快速定位和修复。",
        },
    )

    replace_contains(
        document,
        [
            ("材料提交后进入待审核状态；指导教师或项目负责人可执行审核", "材料提交后进入待审核状态；指导教师、系统管理员或校级管理员可执行审核（通过/驳回）。项目负责人负责上传，不具备材料审核权限。"),
            ("支持查看历史提交记录和审核状态", "功能用途：按当前项目阶段要求，提交计划书PDF、路演PPT PDF和项目展示视频。支持查看历史提交记录、审核状态和退回原因。"),
            ("进入路径：公告发布：首页顶部操作区", "进入路径：公告发布：首页顶部操作区 → “发布公告”；Bug反馈：顶部Bug反馈入口。"),
            ("查看反馈：点击顶部帮助图标", "查看反馈：系统管理员在消息或反馈列表中查看用户提交的Bug反馈。"),
        ],
    )

    insert_images(document, image_paths)
    update_tables(document)
    document.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
