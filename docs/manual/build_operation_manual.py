from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[2]
ASSET_DIR = ROOT / "docs" / "manual" / "assets"
OUT = ROOT / "docs" / "manual" / "中国国际大学生创新大赛管理系统操作手册-详细版.docx"

BLUE = "1F5FEA"
NAVY = "17325C"
LIGHT_BLUE = "EEF5FF"
PALE = "F8FBFF"
GRAY = "64748B"
TEXT = "1E293B"
LINE = "D7E2F2"
AMBER = "FFF7ED"
GREEN = "ECFDF5"


def set_run_font(run, size=None, bold=False, color=None):
    run.font.name = "微软雅黑"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    if size is not None:
        run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_border(cell, color=LINE, size="8"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        node = borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), size)
        node.set(qn("w:space"), "0")
        node.set(qn("w:color"), color)


def cell_margins(cell, top=130, start=160, bottom=130, end=160):
    tc_pr = cell._tc.get_or_add_tcPr()
    mar = tc_pr.first_child_found_in("w:tcMar")
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tc_pr.append(mar)
    for key, value in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def add_p(doc, text="", size=10.5, color=TEXT, bold=False, align=None, before=0, after=4, style=None):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.28
    if align is not None:
        p.alignment = align
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14 if level == 1 else 9)
    p.paragraph_format.space_after = Pt(7 if level == 1 else 5)
    r = p.add_run(text)
    set_run_font(r, size=18 if level == 1 else 14 if level == 2 else 12.5, bold=True, color=NAVY if level == 1 else BLUE)
    return p


def add_note(doc, title, body, fill=PALE):
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell_shading(cell, fill)
    cell_border(cell)
    cell_margins(cell, 180, 220, 180, 220)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    set_run_font(r, 10.5, True, BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(3)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.25
    r2 = p2.add_run(body)
    set_run_font(r2, 10, False, GRAY)
    add_p(doc, "", after=3)


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = 1.22
        r = p.add_run(item)
        set_run_font(r, 10.2, False, TEXT)


def add_table(doc, headers, rows, widths=None, header_fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    if widths is None:
        widths = [Cm(16.2 / len(headers))] * len(headers)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        cell.width = widths[i]
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell_shading(cell, header_fill)
        cell_border(cell)
        cell_margins(cell)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        set_run_font(r, 9.8, True, NAVY)
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].width = widths[i]
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_border(cells[i])
            cell_margins(cells[i])
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i == 0 and len(str(value)) <= 4 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(value))
            set_run_font(r, 9.2, False, TEXT)
    add_p(doc, "", after=5)
    return table


def add_step_table(doc, rows):
    return add_table(
        doc,
        ["步骤", "操作路径/位置", "详细操作办法"],
        rows,
        [Cm(1.6), Cm(4.5), Cm(10.1)],
    )


def add_module_block(doc, title, applies, purpose, path, flow, steps, notes, image_name=None, fig_no=None):
    add_heading(doc, title, 2)
    add_note(doc, "功能用途", purpose)
    add_table(
        doc,
        ["适用角色", "操作路径", "业务流转"],
        [(applies, path, flow)],
        [Cm(3.2), Cm(4.7), Cm(8.3)],
        header_fill="F0F7FF",
    )
    add_step_table(doc, steps)
    if notes:
        add_note(doc, "注意事项", notes, AMBER)
    if image_name and fig_no:
        add_figure(doc, fig_no, title, path, image_name)


def add_figure(doc, fig_no, title, path, image_name):
    image_path = ASSET_DIR / image_name
    add_note(doc, f"截图 {fig_no}：{title}", f"截图位置：{path}。截图已做姓名、账号、邮箱、项目组等信息脱敏处理。", "F8FBFF")
    if image_path.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.add_run().add_picture(str(image_path), width=Cm(16.2))
    add_p(doc, f"图 {fig_no}  {title}", size=9, color=GRAY, align=WD_ALIGN_PARAGRAPH.CENTER, after=8)


def setup_doc(doc):
    section = doc.sections[0]
    section.top_margin = Cm(1.7)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    styles = doc.styles
    styles["Normal"].font.name = "微软雅黑"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
    styles["Normal"].font.size = Pt(10.5)
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = header.add_run("中国国际大学生创新大赛管理系统操作手册（详细版）")
    set_run_font(r, 8.5, False, GRAY)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = footer.add_run("用户单位：南京铁道职业技术学院    支持单位：南京君如玉科技有限公司")
    set_run_font(r2, 8.5, False, GRAY)


def cover(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(78)
    r = p.add_run("中国国际大学生创新大赛\n管理系统操作手册")
    set_run_font(r, 28, True, NAVY)
    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(18)
    r2 = p2.add_run("详细版 · 领导汇报与培训使用")
    set_run_font(r2, 15, True, BLUE)
    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(42)
    r3 = p3.add_run("覆盖：学生端、指导教师端、管理员端")
    set_run_font(r3, 12, False, TEXT)
    p4 = doc.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(72)
    r4 = p4.add_run(f"版本：V2.0    日期：{date.today().isoformat()}")
    set_run_font(r4, 10.5, False, GRAY)
    doc.add_page_break()


def overview_sections(doc):
    add_heading(doc, "一、系统概述", 1)
    add_note(
        doc,
        "建设目标",
        "系统用于支撑中国国际大学生创新大赛项目全过程管理，覆盖账号审核、团队分组、日常汇报、任务工单、项目材料、专家意见、专家评审、现场大屏和资料归档等环节，形成“项目过程可追踪、教师指导可留痕、管理员统筹有依据、现场评审有数据”的闭环管理机制。",
    )
    add_table(
        doc,
        ["模块", "核心用途", "主要使用角色"],
        [
            ("首页概览", "集中展示待办、消息、赛事日程、项目进度和异常提醒。", "学生、教师、管理员"),
            ("时间进度", "展示赛事关键节点、截止时间和阶段安排。", "学生、教师、管理员"),
            ("任务工单", "发布、分配、处理和验收项目推进事项。", "学生、教师、管理员"),
            ("日程汇报", "沉淀每日工作进展，支持教师点评和管理员统计。", "学生、教师、管理员"),
            ("专家意见", "记录专家反馈、整改动作和落实情况。", "学生、教师、管理员"),
            ("项目管理", "配置阶段、提交材料、审核材料和查看项目状态。", "学生、教师、管理员"),
            ("专家评审", "配置评审包、专家席位、评分规则和现场大屏。", "管理员为主，专家参与评分"),
            ("资料归档", "按项目组归档已审核材料和历史版本。", "学生、教师、管理员"),
            ("团队管理", "管理项目组、成员、教师、账号审核和权限。", "管理员为主"),
        ],
        [Cm(3.0), Cm(9.0), Cm(4.2)],
    )
    add_heading(doc, "1.1 系统业务闭环", 2)
    add_table(
        doc,
        ["环节", "系统动作", "责任角色", "结果"],
        [
            ("账号准入", "用户注册后由管理员审核，审核时必须绑定项目组。", "管理员", "账号具备正确身份和数据范围。"),
            ("过程推进", "学生提交日报、处理工单、上传材料。", "学生/项目负责人", "形成过程记录和材料沉淀。"),
            ("教师指导", "教师查看日报、点评、审核材料、督促整改。", "指导教师", "指导过程可追踪。"),
            ("组织统筹", "管理员查看趋势、异常、材料审批、任务进度。", "管理员", "掌握全校或全院项目状态。"),
            ("专家评审", "管理员配置评审包和大屏，专家打分，系统计算最终得分。", "管理员/专家", "生成评审数据和成绩依据。"),
            ("资料归档", "审核通过的材料进入归档，支持后续复盘和留存。", "全角色按权限查看", "项目材料可查询、可追溯。"),
        ],
        [Cm(2.4), Cm(6.0), Cm(3.2), Cm(4.6)],
    )

    add_heading(doc, "1.2 角色职责边界", 2)
    add_table(
        doc,
        ["角色", "主要职责", "不能操作的事项"],
        [
            ("学生/项目负责人", "提交日报、处理工单、上传材料、查看专家意见、配合整改。", "不能审核账号，不能删除账号，不能管理全校数据。"),
            ("指导教师", "查看负责项目组进展、点评日报、审核材料、督促任务和整改。", "不能发布全校公告，不能审核注册账号，不能删除账号。"),
            ("校级管理员", "审核账号、管理项目组、配置项目阶段、组织评审、查看统计。", "不能管理系统管理员账号。"),
            ("系统管理员", "具备全局系统管理权限，负责账号、配置、数据和反馈处理。", "应避免直接替代业务角色完成日常提交。"),
        ],
        [Cm(3.1), Cm(8.0), Cm(5.1)],
    )

    add_heading(doc, "二、账号注册、登录与审核", 1)
    add_module_block(
        doc,
        "2.1 登录系统",
        "所有用户",
        "登录页用于身份认证。系统通过账号、密码和图形验证码确认登录用户身份，避免无关人员进入管理工作台。",
        "系统首页 / 登录页",
        "登录成功后进入与角色对应的工作台；账号未审核、密码错误或验证码错误时不能进入系统。",
        [
            ("1", "登录页", "输入账号名或邮箱。若管理员创建账号，可使用管理员分配的账号登录。"),
            ("2", "密码框", "输入账号密码。首次登录后建议进入个人资料修改密码。"),
            ("3", "验证码区域", "输入右侧图片验证码；看不清时点击验证码图片刷新。"),
            ("4", "登录按钮", "点击“登录”。系统会根据角色自动进入学生、教师或管理员工作台。"),
        ],
        "不要把账号密码交给他人使用。连续输错可能触发登录限制，请稍后再试或联系管理员重置密码。",
        "00-login.png",
        "2-1",
    )
    add_module_block(
        doc,
        "2.2 用户注册与审核",
        "学生、项目负责人、指导教师",
        "注册功能用于新用户申请账号。系统要求邮箱验证码，学生类账号还需填写学院、班级、学号；教师账号需填写所属学院或部门。",
        "登录页 → 注册账号",
        "用户提交注册申请后进入待审核状态；管理员审核时必须分配项目组，审核通过后用户才能登录。",
        [
            ("1", "注册入口", "在登录页点击“注册账号”。"),
            ("2", "邮箱验证码", "填写邮箱后点击“获取验证码”，将邮箱收到的验证码填入注册表。"),
            ("3", "角色选择", "选择“指导教师”“项目负责人”或“团队成员”。不同角色需要填写的信息不同。"),
            ("4", "身份信息", "学生填写学院、专业班级、学号；教师填写所属学院或部门。"),
            ("5", "提交申请", "确认信息无误后提交。提交后不能立即登录，需要管理员审核。"),
            ("6", "管理员审核", "管理员在团队管理中选择项目组并审核通过。未选择项目组时系统不允许通过审核。"),
        ],
        "注册信息应使用真实身份信息，但对外培训材料中的截图必须脱敏。若邮箱收不到验证码，应先检查邮箱地址是否正确，再联系系统管理员确认邮件服务状态。",
        "01-register.png",
        "2-2",
    )


def student_sections(doc):
    doc.add_page_break()
    add_heading(doc, "三、学生端操作办法", 1)
    add_note(
        doc,
        "使用对象",
        "学生端包括项目负责人和团队成员。项目负责人通常承担材料提交、任务统筹和进度推进职责；团队成员按照分工提交日报、处理工单、参与材料完善和意见整改。",
    )
    add_module_block(
        doc,
        "3.1 首页概览",
        "学生、项目负责人",
        "首页用于快速查看本人所在项目组的整体状态，包括待办事项、未读消息、文档审批、任务进度、赛事日程和紧急提醒。",
        "左侧菜单 → 首页概览",
        "首页只展示当前账号有权限查看的数据；学生看到本组数据，项目负责人可看到本组更多任务和材料状态。",
        [
            ("1", "顶部信息栏", "查看当前页面、日期、天气、待办和通知。"),
            ("2", "欢迎横幅", "确认当前登录身份和组织单位是否正确。"),
            ("3", "统计卡片", "查看待分配工单、进行中工单、未读消息、文档待审批等数量。"),
            ("4", "业务进度", "查看团队汇报、任务工单、文档审批、专家评审的完成情况。"),
            ("5", "赛事日程", "查看近期截止事项，提前准备材料、路演或答辩。"),
            ("6", "紧急事项", "优先处理逾期、临近截止或管理员标记的重要事项。"),
        ],
        "首页是工作台入口，不建议只看首页数字；发现异常后应点击对应模块进入详情处理。",
        "02-student-overview.png",
        "3-1",
    )
    add_module_block(
        doc,
        "3.2 日程汇报",
        "学生、项目负责人",
        "日程汇报用于记录每日项目推进情况，帮助教师掌握学生实际工作进展，也为管理员统计项目活跃度提供依据。",
        "左侧菜单 → 日程汇报",
        "学生提交后，本组指导教师可以查看并点评；管理员可以按项目组查看提交率和异常情况。",
        [
            ("1", "选择日期", "默认进入当天日期，也可查看历史日期。新建记录时一般不选择过去日期。"),
            ("2", "填写今日完成", "写清楚当天实际完成的工作，不建议只写“继续推进”。"),
            ("3", "填写问题困难", "如遇技术、材料、人员分工、时间节点问题，应简明说明。"),
            ("4", "填写下一步计划", "写明下一步具体动作、责任人或完成时间。"),
            ("5", "提交日报", "点击提交后形成当日记录，教师端可查看。"),
            ("6", "查看点评", "教师点评或提出修改建议后，学生应及时阅读并在后续工作中落实。"),
        ],
        "日报应围绕项目推进填写，避免无关内容。项目负责人应关注本组成员是否按时提交，必要时线下提醒。",
        "03-student-reports.png",
        "3-2",
    )
    add_module_block(
        doc,
        "3.3 任务工单",
        "学生、项目负责人、团队成员",
        "任务工单用于把项目推进事项拆解成可跟踪任务，适合处理材料修改、路演准备、答辩问题、数据补充等具体工作。",
        "左侧菜单 → 任务工单",
        "管理员、教师或项目负责人发布任务后，处理人完成并提交；需要验收的任务由指定角色确认。",
        [
            ("1", "查看任务列表", "按状态查看全部、待处理、处理中、待验收和已完成任务。"),
            ("2", "阅读任务要求", "打开任务详情，确认任务标题、要求、截止时间和处理人。"),
            ("3", "领取或处理", "属于本人处理的任务，按要求完成并填写处理说明。"),
            ("4", "提交完成", "完成后提交，进入待验收或已完成状态。"),
            ("5", "处理退回", "若任务被退回，应阅读原因，修改后重新提交。"),
        ],
        "系统已取消高中低优先级，所有任务按普通工单处理；实际紧急程度通过标题、截止时间和提醒说明体现。",
        "04-student-board.png",
        "3-3",
    )
    add_module_block(
        doc,
        "3.4 项目管理与材料提交",
        "项目负责人为主，团队成员按权限查看",
        "项目管理用于查看项目阶段和材料要求。项目负责人按阶段上传计划书、路演材料、视频等材料，并等待审核。",
        "左侧菜单 → 项目管理",
        "项目负责人提交材料后，指导教师或管理员审核；审核通过后进入资料归档，审核退回后需修改重传。",
        [
            ("1", "查看当前阶段", "确认当前项目处于计划书、路演、专家评审或其他阶段。"),
            ("2", "阅读材料要求", "查看系统要求上传的材料类型、格式、截止时间和说明。"),
            ("3", "上传文件", "点击上传入口，选择本地文件并提交。文件名建议包含项目名、材料类型和版本。"),
            ("4", "查看审核状态", "提交后关注待审核、已通过、已退回等状态。"),
            ("5", "按意见修改", "如被退回，阅读审核意见，修改文件后重新上传。"),
        ],
        "上传材料应使用最终确认版本，避免频繁提交无效版本。视频、压缩包等大文件应提前检查格式和大小。",
        "05-student-project.png",
        "3-4",
    )
    add_module_block(
        doc,
        "3.5 专家意见与整改落实",
        "学生、项目负责人、团队成员",
        "专家意见用于集中查看专家对项目提出的反馈，帮助团队围绕商业模式、技术路线、路演表达和答辩准备进行整改。",
        "左侧菜单 → 专家意见",
        "管理员或专家录入意见后，本组成员可查看；项目负责人组织整改，教师协助把关。",
        [
            ("1", "查看意见列表", "进入专家意见页面，按项目组或状态查看反馈。"),
            ("2", "阅读反馈摘要", "重点查看专家指出的问题、附件和建议方向。"),
            ("3", "明确整改动作", "将意见拆解为材料修改、数据补充、表达优化或答辩准备等任务。"),
            ("4", "反馈落实情况", "整改完成后在相关任务或材料中体现修改结果。"),
        ],
        "专家意见不是只读信息，应转化为具体整改任务。项目负责人应跟踪每条意见是否落实。",
        "06-student-experts.png",
        "3-5",
    )


def teacher_sections(doc):
    doc.add_page_break()
    add_heading(doc, "四、指导教师端操作办法", 1)
    add_note(
        doc,
        "使用对象",
        "指导教师负责项目过程指导和质量把关。教师端重点不是替学生提交材料，而是查看进度、点评日报、审核材料、督促任务和协助落实专家意见。",
    )
    add_module_block(
        doc,
        "4.1 教师首页概览",
        "指导教师",
        "教师首页集中展示所指导项目组的任务、日报、材料和关键节点，便于教师快速判断哪些项目需要关注。",
        "左侧菜单 → 首页概览",
        "教师只查看与自己绑定项目组相关的数据；如看不到项目组，应联系管理员检查账号分组。",
        [
            ("1", "确认身份", "查看右上角账号和侧边栏角色是否为指导教师。"),
            ("2", "查看待办", "关注待点评日报、待审核材料、待处理工单和近期截止事项。"),
            ("3", "进入详情", "点击首页卡片或左侧菜单进入对应模块处理。"),
            ("4", "处理提醒", "对临近截止或异常项目组及时提醒学生或项目负责人。"),
        ],
        "教师账号不应出现发布全校公告、审核注册账号、删除账号等管理员操作入口。",
        "07-teacher-overview.png",
        "4-1",
    )
    add_module_block(
        doc,
        "4.2 查看日报与发起点评",
        "指导教师",
        "教师通过日程汇报掌握项目组成员每日工作情况，对学生推进情况进行点评、表扬或提出改进建议。",
        "左侧菜单 → 日程汇报",
        "学生提交日报后，教师查看并点评；点评记录学生端可见，管理员也可查看教师指导活跃度。",
        [
            ("1", "选择项目组", "若教师绑定多个项目组，先选择需要查看的项目组。"),
            ("2", "选择日期", "默认查看当天，也可切换历史日期。"),
            ("3", "查看提交情况", "查看已提交、未提交、缺失或异常成员。"),
            ("4", "阅读日报内容", "重点看完成事项、问题困难和下一步计划是否具体。"),
            ("5", "发起点评", "对日报进行点赞、点评或提出修改建议。"),
            ("6", "督促未提交", "对未提交成员通过提醒或线下方式督促。"),
        ],
        "点评应具体到项目问题，避免只写“继续努力”。教师点评会作为过程指导留痕。",
        "08-teacher-reports.png",
        "4-2",
    )
    add_module_block(
        doc,
        "4.3 项目材料审核",
        "指导教师",
        "教师对本组提交的项目材料进行质量审核，确保材料内容、格式、逻辑和阶段要求基本达标。",
        "左侧菜单 → 项目管理",
        "学生提交材料后，教师审核通过则进入归档；退回则学生按意见修改后重新提交。",
        [
            ("1", "进入项目管理", "查看本组当前阶段和材料提交状态。"),
            ("2", "打开材料详情", "查看材料标题、类型、提交人、提交时间和附件。"),
            ("3", "审核通过", "材料符合要求时点击通过，材料进入有效版本。"),
            ("4", "退回修改", "材料不符合要求时填写明确退回意见，说明需要修改的具体内容。"),
            ("5", "跟踪复审", "学生重新提交后，教师再次审核，直到材料通过。"),
        ],
        "退回意见应可执行，例如指出具体页码、具体问题和修改方向。",
        "09-teacher-project.png",
        "4-3",
    )
    add_module_block(
        doc,
        "4.4 资料归档与过程留痕",
        "指导教师",
        "资料归档用于查看本组已通过审核的正式材料和历史版本，便于赛前复盘、答辩准备和后续留存。",
        "左侧菜单 → 资料归档",
        "审核通过的材料会进入归档；教师和学生按项目组权限查看，管理员可按全局范围查看。",
        [
            ("1", "进入资料归档", "查看项目组资料目录和文件列表。"),
            ("2", "筛选资料", "按项目组、阶段、材料类型或状态查找文件。"),
            ("3", "查看文件", "打开预览或下载入口，检查材料版本是否正确。"),
            ("4", "指导完善", "如发现材料仍需补充，应回到项目管理或任务工单发起处理。"),
        ],
        "归档资料应作为正式材料管理，避免用聊天软件传递最终版后系统中无记录。",
        "10-teacher-documents.png",
        "4-4",
    )


def admin_sections(doc):
    doc.add_page_break()
    add_heading(doc, "五、管理员端操作办法", 1)
    add_note(
        doc,
        "使用对象",
        "管理员包括系统管理员和校级管理员。管理员端体现系统治理能力，重点是账号准入、团队组织、项目阶段、专家评审、数据统计和问题反馈处理。",
    )
    add_module_block(
        doc,
        "5.1 管理员首页概览",
        "系统管理员、校级管理员",
        "管理员首页用于查看全校或全院项目总体状态，包括账号审核、任务、消息、资料审批、业务进度和赛事日程。",
        "左侧菜单 → 首页概览",
        "首页提供全局态势，管理员根据异常提醒进入具体模块处理。",
        [
            ("1", "查看概览指标", "查看待审核账号、进行中任务、未读消息、文档待审批等核心指标。"),
            ("2", "查看业务进度", "关注团队汇报、任务工单、文档审批、专家评审等完成情况。"),
            ("3", "查看赛事日程", "关注近期截止事项，提前组织项目组准备。"),
            ("4", "处理紧急事项", "对逾期、未提交、待审核等事项进入对应模块处理。"),
            ("5", "发布公告", "需要全校通知时，通过顶部发布公告入口发布信息。"),
        ],
        "首页数字用于提示风险，最终处理仍应进入对应模块查看明细。",
        "11-admin-overview.png",
        "5-1",
    )
    add_module_block(
        doc,
        "5.2 团队管理、账号审核与分组",
        "系统管理员、校级管理员",
        "团队管理用于维护项目组、成员、教师和账号审核。系统要求注册账号审核时必须分配项目组，保证用户登录后只看到授权范围内的数据。",
        "左侧菜单 → 团队管理",
        "用户注册后进入待审核；管理员审核并绑定项目组后，用户才能登录工作台。",
        [
            ("1", "创建项目组", "根据参赛项目建立项目组，名称建议与项目名称保持一致。"),
            ("2", "维护成员", "将项目负责人、团队成员和指导教师绑定到对应项目组。"),
            ("3", "审核注册申请", "查看待审核账号的姓名、角色、学院/班级/学号或教师部门信息。"),
            ("4", "分配项目组", "审核通过前必须选择项目组；未选择时系统阻止通过。"),
            ("5", "通过或驳回", "信息真实且分组正确则通过；信息不完整或角色错误则驳回。"),
            ("6", "账号维护", "校级管理员和系统管理员可重置密码、调整角色或停用账号；普通学生和教师不显示删除账号入口。"),
        ],
        "账号审核是权限边界的第一道关口。不要让未分组账号通过审核，否则会造成数据范围不明确。",
        "12-admin-team.png",
        "5-2",
    )
    add_module_block(
        doc,
        "5.3 项目阶段与材料要求配置",
        "系统管理员、校级管理员",
        "项目管理用于配置比赛推进阶段、材料要求、提交时间和项目组范围，是材料提交和评审工作的基础。",
        "左侧菜单 → 项目管理",
        "管理员配置阶段后，项目组按要求提交材料；教师或管理员审核后形成有效材料。",
        [
            ("1", "创建阶段", "根据赛事安排创建计划书、路演材料、视频、现场评审等阶段。"),
            ("2", "设置时间", "配置提交开始时间、截止时间和审核要求。新建时间默认不应早于当天。"),
            ("3", "配置材料", "设置每个阶段需要提交的材料类型和说明。"),
            ("4", "选择项目组", "指定该阶段适用的项目组，支持全体或部分项目组。"),
            ("5", "查看提交状态", "按项目组查看未提交、待审核、已通过、已退回状态。"),
            ("6", "处理异常", "对逾期未提交或被退回多次的项目组进行提醒。"),
        ],
        "阶段和材料规则配置后会影响学生端提交入口，应在发布前检查时间和项目组范围。",
        "13-admin-project.png",
        "5-3",
    )
    add_module_block(
        doc,
        "5.4 专家评审、评分规则与现场大屏",
        "系统管理员、校级管理员",
        "专家评审用于组织路演评分、专家席位、评分规则和现场大屏。系统按评审包中的规则计算最终得分，保证后台、大屏和导出结果一致。",
        "左侧菜单 → 专家评审",
        "管理员配置评审包和大屏；专家提交分数；管理员确认并计算最终得分；系统保存结果。",
        [
            ("1", "创建评审包", "选择项目阶段、参评项目组和评审时间。"),
            ("2", "配置专家", "添加或选择评审专家，形成固定匿名席位。"),
            ("3", "设置评分规则", "配置去最高分数量、去最低分数量，系统校验至少保留有效评分。"),
            ("4", "生成大屏链接", "现场投屏前生成链接，用于大屏展示项目、倒计时、匿名席位和分数。"),
            ("5", "配置路演顺序", "不抽签时按项目组配置顺序；抽签时按已设置分组进行组内随机。"),
            ("6", "阶段控制", "依次开始路演、答辩、评分；必要时启动倒计时。"),
            ("7", "专家评分", "专家端提交个人分数后，大屏显示席位状态或单项分数动画。"),
            ("8", "确认最终得分", "所有有效分数完成后，管理员点击确认并计算最终得分，系统按评分规则锁定结果。"),
        ],
        "现场大屏主要用于公开展示，不应暴露真实专家姓名。作废席位仅用于专家离场、设备故障等异常情况。",
        "14-admin-review.png",
        "5-4",
    )
    add_module_block(
        doc,
        "5.5 日程统计与异常分析",
        "系统管理员、校级管理员",
        "日程汇报统计用于查看各项目组日报提交率、趋势变化、教师点评情况和异常项目，帮助管理者及时发现推进风险。",
        "左侧菜单 → 日程汇报",
        "学生提交日报、教师点评后，管理员可按项目组、日期和趋势图查看整体情况。",
        [
            ("1", "选择日期范围", "查看当天、本周或本月数据。当天未到截止时间时系统显示待统计。"),
            ("2", "查看提交率", "关注提交率较低或连续异常项目组。"),
            ("3", "查看趋势图", "趋势图最多显示关键刻度，避免日期挤压；空值处不会误连线。"),
            ("4", "查看教师点评", "关注教师是否及时点评和督促。"),
            ("5", "发起提醒", "对未提交、低活跃或异常项目组进行提醒。"),
        ],
        "趋势图用于判断走势，不应把当天未截止数据误解为 0%。",
        "15-admin-reports.png",
        "5-5",
    )
    add_module_block(
        doc,
        "5.6 任务工单与问题闭环",
        "系统管理员、校级管理员",
        "任务工单用于把管理要求转化为可跟踪事项，包括材料补交、问题整改、赛前准备、账号信息完善等。",
        "左侧菜单 → 任务工单",
        "管理员发布工单，项目负责人、教师或成员处理，完成后进入验收或归档状态。",
        [
            ("1", "发布工单", "填写工单标题、具体要求、截止时间和处理人。"),
            ("2", "选择处理人", "可指定项目负责人、团队成员或指导教师，也可暂不分配进入待分配列表。"),
            ("3", "跟踪状态", "查看待处理、处理中、待验收、已完成等状态。"),
            ("4", "验收结果", "处理完成后管理员或指定审核人确认；不符合要求可退回。"),
            ("5", "沉淀记录", "工单完成记录可作为后续复盘依据。"),
        ],
        "工单标题应具体，避免只写“处理一下”。系统中不再使用高中低优先级，紧急事项通过截止时间和标题说明体现。",
        "16-admin-board.png",
        "5-6",
    )


def closing_sections(doc):
    doc.add_page_break()
    add_heading(doc, "六、权限说明表", 1)
    add_table(
        doc,
        ["功能模块", "学生/项目负责人", "指导教师", "校级管理员/系统管理员"],
        [
            ("账号注册", "可提交注册申请", "可提交注册申请", "审核账号并分配项目组"),
            ("首页概览", "查看本组待办和进度", "查看负责项目组状态", "查看全局数据和异常"),
            ("日程汇报", "提交日报、查看点评", "查看日报、发起点评", "查看统计、趋势和异常"),
            ("任务工单", "处理本人或本组任务", "分配或确认本组任务", "发布、分配、验收和跟踪"),
            ("项目管理", "提交或查看项目材料", "审核本组材料", "配置阶段、查看全局材料"),
            ("专家意见", "查看并落实整改", "指导学生整改", "录入、分配和跟踪意见"),
            ("专家评审", "一般只查看相关结果", "按权限查看相关项目", "配置评审包、评分规则和大屏"),
            ("团队管理", "查看本组成员", "查看负责项目组成员", "管理账号、分组和权限"),
            ("资料归档", "查看本组归档资料", "查看本组归档资料", "查看全局归档资料"),
        ],
        [Cm(3.0), Cm(4.2), Cm(4.2), Cm(4.8)],
    )
    add_heading(doc, "七、常见问题与处理办法", 1)
    add_table(
        doc,
        ["问题", "原因", "处理办法"],
        [
            ("注册后无法登录", "账号仍处于待审核状态，或未分配项目组。", "联系管理员审核账号；管理员必须选择项目组后才能通过。"),
            ("验证码看不清", "图形验证码刷新或加载不完整。", "点击验证码图片刷新，再重新输入。"),
            ("看不到某个菜单", "角色权限不同或账号分组不正确。", "联系管理员核查账号角色和项目组绑定。"),
            ("材料被退回", "材料内容、格式或阶段要求不符合。", "查看退回意见，修改后重新上传。"),
            ("日报趋势图当天异常", "当天未到截止时间，数据仍在统计中。", "等待截止后查看，或查看历史完整日期。"),
            ("大屏不能进入下一项目", "当前轮次、阶段或项目顺序配置未完成。", "管理员检查评审包、分组、项目顺序和阶段状态。"),
            ("专家分数不一致", "未按评审包评分规则统一计算。", "以管理员点击“确认并计算最终得分”后的后端结果为准。"),
            ("需要反馈系统问题", "页面显示异常、权限异常或流程不合理。", "通过顶部帮助/反馈入口提交 Bug 反馈，系统管理员处理。"),
        ],
        [Cm(4.0), Cm(5.6), Cm(6.6)],
    )
    add_note(
        doc,
        "对外发布提醒",
        "本手册截图已做脱敏处理。正式对外发布时，应再次检查是否存在真实姓名、账号、邮箱、手机号、学号、项目组真实名称、专家姓名等敏感信息。",
        AMBER,
    )


def create_document():
    doc = Document()
    setup_doc(doc)
    cover(doc)
    overview_sections(doc)
    student_sections(doc)
    teacher_sections(doc)
    admin_sections(doc)
    closing_sections(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(create_document())
