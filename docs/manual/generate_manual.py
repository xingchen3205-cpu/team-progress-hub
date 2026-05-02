#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成《中国国际大学生创新大赛管理系统操作手册》Word文档
使用说明：python3 docs/manual/generate_manual.py
"""

import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ============================================================
# 配置
# ============================================================
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
ASSETS_DIR = os.path.join(BASE_DIR, "assets")
OUTPUT_PATH = os.path.join(BASE_DIR, "中国国际大学生创新大赛管理系统操作手册-正式版.docx")

# 截图映射: (图号, 文件名, 图注)
FIGURES = {
    "4-1": ("00-login.png", "图 4-1  系统登录页面"),
    "4-2": ("01-register.png", "图 4-2  账号注册页面"),
    "5-1": ("02-student-overview.png", "图 5-1  学生端首页概览"),
    "5-2": ("03-student-reports.png", "图 5-2  学生端日程汇报页面"),
    "5-3": ("04-student-board.png", "图 5-3  学生端任务工单页面"),
    "5-4": ("05-student-project.png", "图 5-4  学生端项目管理页面"),
    "5-5": ("06-student-experts.png", "图 5-5  学生端专家意见页面"),
    "6-1": ("07-teacher-overview.png", "图 6-1  指导教师端首页概览"),
    "6-2": ("08-teacher-reports.png", "图 6-2  指导教师端日程汇报与点评"),
    "6-3": ("09-teacher-project.png", "图 6-3  指导教师端项目管理页面"),
    "6-4": ("10-teacher-documents.png", "图 6-4  指导教师端资料归档页面"),
    "7-1": ("11-admin-overview.png", "图 7-1  管理员端首页概览"),
    "7-2": ("12-admin-team.png", "图 7-2  管理员端团队管理与账号审核"),
    "7-3": ("13-admin-project.png", "图 7-3  管理员端项目阶段与材料管理"),
    "7-4": ("14-admin-review.png", "图 7-4  管理员端专家评审组织"),
    "7-5": ("15-admin-reports.png", "图 7-5  管理员端日程汇报统计"),
    "7-6": ("16-admin-board.png", "图 7-6  管理员端全校任务台账"),
}


def set_cell_border(cell, **kwargs):
    """为单元格设置边框（简单实现，部分场景用）"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = "w:{}".format(edge)
            element = OxmlElement(tag)
            element.set(qn("w:val"), edge_data.get("val", "single"))
            element.set(qn("w:sz"), str(edge_data.get("sz", 4)))
            element.set(qn("w:space"), "0")
            element.set(qn("w:color"), edge_data.get("color", "000000"))
            tcBorders.append(element)
    tcPr.append(tcBorders)


def set_run_font(run, font_name="宋体", size=10.5, bold=False, color=None):
    font = run.font
    font.name = font_name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font_name)
    font.size = Pt(size)
    font.bold = bold
    if color:
        font.color.rgb = RGBColor(*color)


def add_heading_custom(doc, text, level=1):
    """添加自定义标题，使用中文字体"""
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
    style_name = style_map.get(level, "Heading 1")
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        if level == 1:
            set_run_font(run, "黑体", 16, bold=True, color=(0x1A, 0x6F, 0xD4))
        elif level == 2:
            set_run_font(run, "黑体", 14, bold=True)
        elif level == 3:
            set_run_font(run, "黑体", 12, bold=True)
        else:
            set_run_font(run, "黑体", 11, bold=True)
    p.paragraph_format.space_before = Pt(12 if level > 1 else 18)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.5
    return p


def add_paragraph_custom(doc, text, indent=True, bold=False, size=10.5, align=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=0.0):
    p = doc.add_paragraph()
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    set_run_font(run, "宋体", size, bold=bold)
    p.alignment = align
    return p


def add_bullet_paragraph(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet" if level == 0 else "List Bullet 2")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
    run = p.add_run(text)
    set_run_font(run, "宋体", 10.5)
    return p


def add_numbered_paragraph(doc, text, level=0):
    p = doc.add_paragraph(style="List Number" if level == 0 else "List Number 2")
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
    run = p.add_run(text)
    set_run_font(run, "宋体", 10.5)
    return p


def add_figure(doc, fig_key):
    info = FIGURES.get(fig_key)
    if not info:
        return
    filename, caption = info
    filepath = os.path.join(ASSETS_DIR, filename)
    if os.path.exists(filepath):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(filepath, width=Inches(5.8))
    else:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"【截图待补充：{caption}】")
        set_run_font(run, "宋体", 10, bold=False, color=(0x99, 0x99, 0x99))

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    set_run_font(run, "宋体", 10, bold=False, color=(0x66, 0x66, 0x66))
    p.paragraph_format.space_after = Pt(8)


def add_placeholder_figure(doc, caption_text, note=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"【截图待补充：{caption_text}】")
    set_run_font(run, "宋体", 10, bold=False, color=(0x99, 0x99, 0x99))
    if note:
        run = p.add_run(f"\n{note}")
        set_run_font(run, "宋体", 9, bold=False, color=(0xAA, 0xAA, 0xAA))
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption_text)
    set_run_font(run, "宋体", 10, bold=False, color=(0x66, 0x66, 0x66))
    p.paragraph_format.space_after = Pt(8)


def add_info_block(doc, title, lines):
    """添加一个带标题的信息块（功能用途、进入路径等）"""
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(f"{title}：")
    set_run_font(run, "黑体", 10.5, bold=True)
    if isinstance(lines, str):
        run = p.add_run(lines)
        set_run_font(run, "宋体", 10.5)
    else:
        for i, line in enumerate(lines):
            if i == 0:
                run = p.add_run(line)
                set_run_font(run, "宋体", 10.5)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing = 1.5
                p.paragraph_format.space_after = Pt(3)
                p.paragraph_format.left_indent = Cm(0.5)
                run = p.add_run(line)
                set_run_font(run, "宋体", 10.5)


def add_step_table(doc, steps):
    """添加操作步骤表格 steps: [(步骤号, 操作位置, 操作说明), ...]"""
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    table.columns[0].width = Cm(1.8)
    table.columns[1].width = Cm(3.5)
    table.columns[2].width = Cm(9.7)

    hdr_cells = table.rows[0].cells
    headers = ["步骤", "操作位置", "操作说明"]
    for i, h in enumerate(headers):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, "黑体", 10.5, bold=True)
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for step in steps:
        row_cells = table.add_row().cells
        for i, text in enumerate(step):
            p = row_cells[i].paragraphs[0]
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, "宋体", 10.5)
            row_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph()  # 表格后空一行


def add_permission_table(doc, rows):
    """添加权限表格 rows: [(操作, admin, school_admin, teacher, leader, member, expert), ...]"""
    table = doc.add_table(rows=1, cols=7)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    widths = [4.0, 2.2, 2.2, 2.2, 2.2, 2.2, 2.2]
    for i, w in enumerate(widths):
        table.columns[i].width = Cm(w)

    hdr = ["操作/权限", "系统管理员", "校级管理员", "指导教师", "项目负责人", "团队成员", "评审专家"]
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(hdr):
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, "黑体", 10, bold=True)
        hdr_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    for row in rows:
        row_cells = table.add_row().cells
        for i, text in enumerate(row):
            p = row_cells[i].paragraphs[0]
            if i == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            set_run_font(run, "宋体", 10)
            row_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER

    doc.add_paragraph()


# ============================================================
# 文档生成主逻辑
# ============================================================
def build_document():
    doc = Document()

    # 页面设置 A4
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    # 修改默认样式字体
    style = doc.styles["Normal"]
    style.font.name = "宋体"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    style.font.size = Pt(10.5)

    # ============================================================
    # 封面
    # ============================================================
    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("中国国际大学生创新大赛管理系统")
    set_run_font(run, "黑体", 28, bold=True, color=(0x1A, 0x6F, 0xD4))
    p.paragraph_format.space_after = Pt(12)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("操 作 手 册")
    set_run_font(run, "黑体", 28, bold=True, color=(0x1A, 0x6F, 0xD4))
    p.paragraph_format.space_after = Pt(60)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("学生端 · 指导教师端 · 管理员端 · 评审专家端")
    set_run_font(run, "宋体", 14, bold=False, color=(0x33, 0x33, 0x33))
    p.paragraph_format.space_after = Pt(80)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("版本：V2.0\n日期：2026年5月1日\n编制单位：南京铁道职业技术学院")
    set_run_font(run, "宋体", 12, bold=False, color=(0x66, 0x66, 0x66))

    doc.add_page_break()

    # ============================================================
    # 第一章 系统概述
    # ============================================================
    add_heading_custom(doc, "第一章 系统概述", level=1)

    add_heading_custom(doc, "1.1 编写目的", level=2)
    add_paragraph_custom(doc,
        "本手册面向参赛学生、指导教师、校级管理员、系统管理员及评审专家，旨在提供系统各功能模块的标准化操作流程说明。手册以\"在哪里点、填什么、提交后谁处理\"为主线，帮助用户快速上手并规范使用系统，同时为学校领导汇报和系统正式发布提供配套文档支撑。",
        first_line_indent=0.74)

    add_heading_custom(doc, "1.2 适用对象", level=2)
    add_paragraph_custom(doc,
        "本手册适用于以下六类系统角色：系统管理员、校级管理员、指导教师、项目负责人、团队成员和评审专家。不同角色在系统中的可见菜单、可操作功能和数据范围存在差异，具体以第三章《角色权限说明》为准。",
        first_line_indent=0.74)

    add_heading_custom(doc, "1.3 运行环境", level=2)
    add_bullet_paragraph(doc, "浏览器：推荐使用 Chrome、Edge、Firefox、Safari 等现代浏览器（建议保持最新版本）。")
    add_bullet_paragraph(doc, "网络：校园网或公网均可访问，建议保持稳定的网络连接以确保材料上传和视频预览顺畅。")
    add_bullet_paragraph(doc, "分辨率：建议显示器分辨率不低于 1366×768，以获得最佳侧边栏和表格展示效果。")

    add_heading_custom(doc, "1.4 术语说明", level=2)
    terms = [
        ("项目组（Team Group）", "参赛项目的基本组织单元，由项目负责人、团队成员和指导教师组成。账号必须绑定项目组后才能正常使用系统功能。"),
        ("工单（Task / Board）", "任务管理的基本单位，包含提报、分配、接取、处理、验收、归档六个状态，形成闭环管理。"),
        ("项目阶段（Project Stage）", "管理员定义的赛事环节，分为\"网络评审\"和\"项目路演\"两种类型，用于控制材料提交要求和评审方式。"),
        ("评审包（Review Package）", "专家评审的基本单元，包含目标项目、评审专家、评分规则和时间窗口等配置信息。"),
        ("路演（Roadshow）", "现场答辩评审模式，不要求提前上传材料，通过大屏控制台完成抽签、路演、答辩、评分和揭晓分数全流程。"),
        ("网评（Network Review）", "线上材料评审模式，专家通过系统在线预览项目提交的PDF计划书、路演PPT和视频，并直接输入总分完成评分。"),
        ("日程汇报（Daily Report）", "学生按日期提交的工作日报，包含当日完成工作、遇到的问题和下一步计划，支持指导教师点评。"),
    ]
    for term, desc in terms:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"{term}：")
        set_run_font(run, "黑体", 10.5, bold=True)
        run = p.add_run(desc)
        set_run_font(run, "宋体", 10.5)

    doc.add_page_break()

    # ============================================================
    # 第二章 业务流程总览
    # ============================================================
    add_heading_custom(doc, "第二章 业务流程总览", level=1)

    add_heading_custom(doc, "2.1 大赛管理全生命周期", level=2)
    add_paragraph_custom(doc,
        "系统围绕中国国际大学生创新大赛的组织管理，形成\"账号准入—项目筹备—材料提交—专家评审—结果归档\"五大阶段闭环。各阶段的核心参与角色和系统支撑功能如下：",
        first_line_indent=0.74)
    add_numbered_paragraph(doc, "账号准入阶段：学生/教师自助注册 → 管理员审核并分配项目组 → 审核通过后登录。")
    add_numbered_paragraph(doc, "项目筹备阶段：管理员创建项目阶段、配置材料要求和时间窗口；学生组建团队、分配角色。")
    add_numbered_paragraph(doc, "材料提交阶段：学生按阶段要求上传计划书PDF、路演PPT PDF、项目视频；指导教师/项目负责人审核；通过后进入资料归档。")
    add_numbered_paragraph(doc, "专家评审阶段：管理员创建评审包、分配专家、设置评分规则；专家在线评分或现场路演评分；系统自动计算最终得分。")
    add_numbered_paragraph(doc, "结果归档阶段：评分结果、审核通过的材料、日报记录等统一归档，支持按项目组查询和导出。")

    add_heading_custom(doc, "2.2 账号生命周期", level=2)
    add_paragraph_custom(doc,
        "系统中所有学生端和教师端账号均需经过严格的准入流程，具体流转如下：",
        first_line_indent=0.74)
    add_step_table(doc, [
        ("1", "登录页", "点击\"注册账号\"，进入自助注册页面。"),
        ("2", "注册表单", "填写姓名、账号名、邮箱、验证码、密码，选择角色（项目负责人/团队成员/指导教师），补充学院、班级、学号/部门等信息。"),
        ("3", "提交注册", "提交后账号状态为\"待审核\"，此时无法登录系统。"),
        ("4", "管理员审核", "系统管理员或校级管理员在\"团队管理\"中查看待审核账号，必须为账号分配项目组后方可审核通过。"),
        ("5", "审核通过", "账号状态变为\"已通过\"，用户可使用账号名/密码登录系统。"),
        ("6", "日常使用", "登录后可访问对应角色权限范围内的功能模块。"),
    ])
    add_paragraph_custom(doc,
        "重要规则：注册账号必须分配项目组后才能审核通过。未分配项目组的待审核账号，管理员无法执行通过操作。",
        bold=True, first_line_indent=0.74)

    add_heading_custom(doc, "2.3 材料生命周期", level=2)
    add_step_table(doc, [
        ("1", "管理员配置", "管理员创建项目阶段，设置阶段类型（网络评审/项目路演）、材料要求、开放项目组和时间窗口。"),
        ("2", "学生上传", "学生在\"项目管理\"中查看当前开放阶段，按材料要求上传文件。"),
        ("3", "审核流转", "材料提交后进入待审核状态；指导教师或项目负责人可执行审核（通过/退回修改）。"),
        ("4", "归档入库", "审核通过的材料自动进入\"资料归档\"，形成历史版本链，支持后续评审调用和下载。"),
    ])

    add_heading_custom(doc, "2.4 评审生命周期", level=2)
    add_step_table(doc, [
        ("1", "创建评审包", "管理员在\"专家评审\"中创建评审包，选择关联的项目阶段和参评项目组。"),
        ("2", "分配专家", "从专家库中选择评审专家，一个评审包可分配多名专家。"),
        ("3", "设置规则", "配置评分规则，包括是否去掉最高分、去掉最低分，以及评审开始时间和截止时间。"),
        ("4", "专家评分", "网评模式下专家在线预览材料后直接输入总分（保留两位小数）；路演模式下专家在大屏评分环节实时打分。"),
        ("5", "计算结果", "系统根据评分规则和专家提交情况，自动计算最终得分；路演模式下支持现场揭晓分数。"),
        ("6", "结果归档", "评分结果与评审记录同步保存，管理员和学生可在系统中查看。"),
    ])

    doc.add_page_break()

    # ============================================================
    # 第三章 角色权限说明
    # ============================================================
    add_heading_custom(doc, "第三章 角色权限说明", level=1)

    add_heading_custom(doc, "3.1 角色对照表", level=2)
    add_paragraph_custom(doc, "系统共设置六个角色，各角色的职责范围和可访问菜单如下：", first_line_indent=0.74)

    role_table = [
        ("系统管理员", "全局管理", "全校管理", "首页概览、时间进度、任务中心、日程汇报、专家意见、项目管理、专家评审、资料归档、团队管理、AI助手、个人信息"),
        ("校级管理员", "全局管理", "全校管理", "首页概览、时间进度、任务中心、日程汇报、专家意见、项目管理、专家评审、资料归档、团队管理、AI助手、个人信息"),
        ("指导教师", "项目组指导", "所在项目组", "首页概览、时间进度、任务中心、训练中心、日程汇报、专家意见、资料归档、项目管理、团队管理、AI助手、个人信息"),
        ("项目负责人", "项目执行", "所在项目组", "首页概览、时间进度、任务中心、训练中心、日程汇报、专家意见、资料归档、项目管理、团队管理、AI助手、个人信息"),
        ("团队成员", "项目执行", "所在项目组", "首页概览、时间进度、任务中心、训练中心、日程汇报、专家意见、资料归档、项目管理、团队管理、AI助手、个人信息"),
        ("评审专家", "独立评审", "专家评审任务", "专家评审、个人信息"),
    ]
    table = doc.add_table(rows=1, cols=4)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.columns[0].width = Cm(2.5)
    table.columns[1].width = Cm(2.5)
    table.columns[2].width = Cm(3.0)
    table.columns[3].width = Cm(7.0)
    hdr = ["角色名称", "职责定位", "数据范围", "可访问菜单"]
    for i, h in enumerate(hdr):
        p = table.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        set_run_font(run, "黑体", 10, bold=True)
        table.rows[0].cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for row in role_table:
        row_cells = table.add_row().cells
        for i, text in enumerate(row):
            p = row_cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i < 3 else WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(text)
            set_run_font(run, "宋体", 10)
            row_cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    doc.add_paragraph()

    add_heading_custom(doc, "3.2 权限边界清单", level=2)
    add_paragraph_custom(doc, "以下列出系统中关键操作的权限边界，用户在进行操作前应确认自身角色是否具备对应权限：", first_line_indent=0.74)

    add_permission_table(doc, [
        ("删除账号", "✓", "✓", "✗", "✗", "✗", "✗"),
        ("审核注册申请", "✓", "✓", "✗", "✗", "✗", "✗"),
        ("重置他人密码", "✓", "✓", "✓", "✗", "✗", "✗"),
        ("创建项目阶段", "✓", "✓", "✓", "✗", "✗", "✗"),
        ("发布全校公告", "✓", "✓", "✗", "✗", "✗", "✗"),
        ("创建/编辑工单", "✓", "✓", "✓", "✓", "✗", "✗"),
        ("删除工单", "✓", "✓", "✓", "✗", "✗", "✗"),
        ("审核项目材料", "✓", "✓", "✓", "✓", "✗", "✗"),
        ("提交日程汇报", "✗", "✗", "✗", "✓", "✓", "✗"),
        ("专家打分", "✗", "✗", "✗", "✗", "✗", "✓"),
        ("配置专家评审包", "✓", "✓", "✗", "✗", "✗", "✗"),
        ("控制现场大屏", "✓", "✓", "✗", "✗", "✗", "✗"),
    ])

    add_heading_custom(doc, "3.3 重要权限规则说明", level=2)
    add_numbered_paragraph(doc, "账号删除：仅系统管理员和校级管理员具备删除账号权限。学生、指导教师、项目负责人、团队成员均不显示删除账号入口。系统管理员可删除任何非系统管理员账号；校级管理员可删除非校级管理员账号。")
    add_numbered_paragraph(doc, "账号审核：账号审核和项目组分配由校级管理员或系统管理员完成。学生、指导教师、项目负责人、团队成员均不能审核他人账号，也无法自行通过审核。")
    add_numbered_paragraph(doc, "项目组分配：注册账号必须分配项目组后才能审核通过。这是系统强制规则，管理员在审核界面若不选择项目组，则无法完成通过操作。")
    add_numbered_paragraph(doc, "数据隔离：除系统管理员和校级管理员可查看全校数据外，指导教师、项目负责人和团队成员仅能查看所在项目组的数据；评审专家仅能看到自己被分配的评审任务。")
    add_numbered_paragraph(doc, "菜单可见性：若发现左侧菜单缺少某项功能，请先确认当前账号角色是否具备该权限，以及账号是否已完成审核并绑定项目组。")

    doc.add_page_break()

    # ============================================================
    # 第四章 登录与注册
    # ============================================================
    add_heading_custom(doc, "第四章 登录与注册", level=1)

    add_heading_custom(doc, "4.1 系统登录", level=2)
    add_info_block(doc, "功能用途", "用户通过账号名、密码和图形验证码进入系统工作台。")
    add_info_block(doc, "进入路径", "浏览器访问系统地址 → 进入登录页。")
    add_info_block(doc, "前置条件", "账号已通过管理员审核且已绑定项目组；评审专家账号由管理员直接创建，无需自助注册。")
    add_info_block(doc, "操作步骤", [
        "在登录页输入账号名（注册时填写的账号名或邮箱均可）。",
        "输入密码。",
        "输入图形验证码；如看不清，可点击验证码图片刷新后重新输入。",
        "点击\"登录\"按钮，进入系统首页概览。",
    ])
    add_info_block(doc, "注意事项", [
        "若连续多次输入错误密码，账号可能被临时锁定，请联系管理员重置密码。",
        "登录状态会在一定时间后失效，若页面提示\"登录状态已失效，正在返回登录页\"，请重新登录。",
        "系统支持通过\"忘记密码\"入口自助重置密码，重置链接将发送至注册邮箱。",
    ])
    add_figure(doc, "4-1")

    add_heading_custom(doc, "4.2 账号注册", level=2)
    add_info_block(doc, "功能用途", "面向指导教师、项目负责人和团队成员提供自助注册通道。注册后账号处于\"待审核\"状态，需等待管理员审核并分配项目组后方可登录。")
    add_info_block(doc, "进入路径", "登录页 → 点击\"注册账号\" → 进入注册表单。")
    add_info_block(doc, "前置条件", "拥有可正常接收邮件的邮箱地址；该邮箱未在系统中注册过。")
    add_info_block(doc, "操作步骤", [
        "填写真实姓名。",
        "填写账号名（登录时使用，建议与姓名或学号关联，便于记忆）。",
        "填写邮箱地址，点击\"获取验证码\"，系统将验证码发送至该邮箱。",
        "在邮箱中查收验证码（有效期有限，请尽快填写），将验证码填入对应字段。",
        "设置密码（至少6位）。",
        "选择角色：指导教师、项目负责人或团队成员。",
        "填写所属学院；项目负责人和团队成员还须填写专业班级和学号；指导教师须填写所属学院或部门。",
        "阅读并确认相关条款后，点击\"提交注册\"。",
    ])
    add_info_block(doc, "提交后流转", "注册成功后，页面提示\"注册成功，等待审核\"。账号状态为\"待审核\"，此时无法登录。系统管理员或校级管理员在\"团队管理\"中收到待审核提醒，审核时必须为该账号分配项目组，分配后点击\"通过\"，账号状态变为\"已通过\"，用户方可正常登录。")
    add_info_block(doc, "注意事项", [
        "一个邮箱只能注册一个账号。若提示\"账号名或邮箱已存在\"，请更换后再试。",
        "项目负责人和团队成员必须填写学院、专业班级和学号，否则无法提交注册。",
        "指导教师必须填写所属学院或部门，否则无法提交注册。",
        "注册时未选择项目组，项目组由管理员在审核环节统一分配。",
    ])
    add_figure(doc, "4-2")

    add_heading_custom(doc, "4.3 密码找回与重置", level=2)
    add_paragraph_custom(doc,
        "用户忘记密码时，可在登录页点击\"忘记密码\"，输入注册时使用的邮箱地址，系统将发送密码重置链接。点击链接后按页面提示设置新密码即可。",
        first_line_indent=0.74)
    add_paragraph_custom(doc,
        "管理员也可在\"团队管理\"中为指定账号直接重置密码，重置后的默认密码由系统规则生成，用户首次登录后建议立即修改为个人密码。",
        first_line_indent=0.74)

    add_heading_custom(doc, "4.4 登录状态失效处理", level=2)
    add_paragraph_custom(doc,
        "当系统检测到登录状态过期或用户在另一设备登录时，页面将提示\"登录状态已失效，正在返回登录页\"。此时请重新输入账号密码登录，未保存的数据可能丢失，建议在操作关键内容时及时保存。",
        first_line_indent=0.74)

    doc.add_page_break()

    # ============================================================
    # 第五章 学生端操作办法
    # ============================================================
    add_heading_custom(doc, "第五章 学生端操作办法", level=1)
    add_paragraph_custom(doc,
        "学生端包括项目负责人和团队成员两类角色。两者的系统界面基本一致，但在部分功能的操作权限上存在差异。本章以通用操作描述为主，并在涉及权限差异处单独标注说明。",
        first_line_indent=0.74)

    # 5.1 首页概览
    add_heading_custom(doc, "5.1 首页概览", level=2)
    add_info_block(doc, "功能用途", "查看项目组基本信息、待办提醒、进度节点、关键统计数据和最新通知。首页是登录后的默认落地页，提供各核心模块的快捷入口。")
    add_info_block(doc, "进入路径", "登录成功后自动进入；或点击左侧菜单\"首页概览\"。")
    add_info_block(doc, "前置条件", "账号已审核通过且已绑定项目组。")
    add_info_block(doc, "操作步骤", [
        "查看顶部通知区：未读消息、待办事项和系统公告以徽标或卡片形式展示，点击可跳转对应模块。",
        "查看进度节点卡片：首页中央区域展示当前项目阶段倒计时、材料提交截止时间和待办任务数量。",
        "查看关键统计数据：包括本组日报提交率、任务完成率、材料审核状态等汇总指标。",
        "使用快捷入口：根据首页提醒，一键进入日程汇报、项目管理或任务工单处理页面。",
    ])
    add_info_block(doc, "提交后流转", "首页数据为只读展示，不涉及提交流转。操作后的状态变更会实时反映在首页统计卡片中。")
    add_info_block(doc, "注意事项", [
        "若首页显示\"未加入项目组\"，说明管理员尚未完成项目组分配，请联系管理员处理。",
        "项目负责人和团队成员在首页看到的数据范围相同，但项目负责人具备更多管理操作按钮（如创建工单）。",
    ])
    add_figure(doc, "5-1")

    # 5.2 时间进度
    add_heading_custom(doc, "5.2 时间进度", level=2)
    add_info_block(doc, "功能用途", "集中查看比赛、答辩、材料提交等关键时间节点，帮助学生掌握赛事节奏。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"时间进度\"。")
    add_info_block(doc, "前置条件", "无特殊前置条件。")
    add_info_block(doc, "操作步骤", [
        "进入时间进度页面，查看事件列表。",
        "按日期远近浏览事件，关注最近待办的关键节点。",
        "点击事件卡片可查看事件详情（时间、地点、说明）。",
    ])
    add_info_block(doc, "注意事项", [
        "学生端对时间进度为只读权限，无法创建或编辑事件。",
        "管理员和指导教师具备编辑权限，可新增、修改或删除时间事件。",
    ])
    add_placeholder_figure(doc, "图 5-2  时间进度页面", "建议截取：左侧菜单\"时间进度\"入口、事件列表、最近待办节点。")

    # 5.3 日程汇报
    add_heading_custom(doc, "5.3 日程汇报", level=2)
    add_info_block(doc, "功能用途", "按日期填写个人日报，记录当天完成的工作内容、遇到的问题和下一步计划。支持查看历史汇报和指导教师点评。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"日程汇报\"。")
    add_info_block(doc, "前置条件", "无特殊前置条件，任何时间均可提交。")
    add_info_block(doc, "操作步骤", [
        "在日程汇报页面，选择要填写的日期（默认为当天）。",
        "在汇报表单中填写：今日完成工作、遇到的问题、下一步计划。",
        "确认内容无误后，点击\"保存\"或\"提交\"。",
        "提交成功后，可在页面下方查看该日期的历史汇报记录。",
        "若指导教师已对本条汇报进行点评，可在记录下方查看点评内容（包括点赞、改进建议或文字评论）。",
    ])
    add_info_block(doc, "提交后流转", "提交成功后，汇报记录即时保存。指导教师可在其\"日程汇报\"页面查看本组全部成员的汇报并进行点评。")
    add_info_block(doc, "注意事项", [
        "团队成员只能查看和编辑自己的汇报记录。",
        "项目负责人可查看本组全部成员的汇报记录，但只能编辑自己的汇报。",
        "建议每日按时提交，保持汇报连续性，便于指导教师跟踪项目进展。",
        "若当日未提交，指导教师可通过系统发送督促提醒。",
    ])
    add_figure(doc, "5-2")

    # 5.4 任务工单
    add_heading_custom(doc, "5.4 任务工单", level=2)
    add_info_block(doc, "功能用途", "以工单形式管理项目执行过程中的各类任务，覆盖提报、分配、接取、处理、验收和归档六个状态，实现任务闭环管理。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"任务中心\"。")
    add_info_block(doc, "前置条件", "无特殊前置条件。")
    add_info_block(doc, "操作步骤", [
        "查看工单列表：页面以看板形式展示全部工单，按状态分为待分配/待接取、处理中、待验收、已完成、已归档等列。",
        "接取任务：在\"待分配/待接取\"列中找到分配给本人的工单，点击接取，状态变为\"处理中\"。",
        "处理任务：按要求完成任务后，在处理说明中填写完成情况，可上传附件作为佐证。",
        "提交验收：点击\"提交验收\"，工单状态变为\"待验收\"，等待项目负责人、指导教师或管理员验收。",
        "查看结果：验收通过后状态变为\"已完成\"；若被退回，状态回到\"处理中\"，需根据退回意见修改后重新提交。",
    ])
    add_info_block(doc, "提交后流转", "todo → doing（接取）→ review（提交验收）→ done（验收通过）→ archived（归档）。")
    add_info_block(doc, "注意事项", [
        "团队成员只能接取分配给本人的工单，无权创建或删除工单。",
        "项目负责人可以创建和编辑工单，但不能删除工单（删除权限仅管理员和指导教师具备）。",
        "工单逾期未完成时，系统可能向相关人员发送提醒。",
    ])
    add_figure(doc, "5-3")

    # 5.5 项目管理（材料提交）
    add_heading_custom(doc, "5.5 项目管理（材料提交）", level=2)
    add_info_block(doc, "功能用途", "按当前项目阶段要求，提交计划书PDF、路演PPT PDF和项目展示视频。支持查看历史提交记录和审核状态。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"项目管理\"。")
    add_info_block(doc, "前置条件", "管理员已创建项目阶段，且当前阶段处于开放窗口期（开始时间已到且截止时间未到）。")
    add_info_block(doc, "操作步骤", [
        "进入项目管理页面，查看当前开放的项目阶段卡片。",
        "阅读阶段说明和材料要求：系统会明确列出需要提交的材料类型和格式限制。",
        "点击对应材料类型的\"上传\"按钮，选择本地文件。",
        "等待文件上传完成，确认文件名称和大小无误后提交。",
        "提交后，材料状态显示为\"待审核\"，等待指导教师或项目负责人审核。",
    ])
    add_info_block(doc, "提交后流转", "pending（待审核）→ leader_approved / teacher_approved（负责人/教师通过）→ approved（最终通过，进入资料归档）或 revision（退回修改）。")
    add_info_block(doc, "注意事项", [
        "计划书和路演PPT仅支持PDF格式上传，请提前将Word或PPT文件导出为PDF。",
        "视频支持 mp4、mov、avi 格式，网络评审阶段视频大小不超过20MB。",
        "若材料被退回，请仔细阅读退回意见，修改后重新上传。",
        "每个阶段对项目组有开放范围限制，若提示\"当前阶段未对您的项目组开放\"，请联系管理员确认项目组配置。",
    ])
    add_figure(doc, "5-4")

    # 5.6 专家意见查看
    add_heading_custom(doc, "5.6 专家意见查看", level=2)
    add_info_block(doc, "功能用途", "按时间倒序查看专家对本组项目提出的反馈意见、整改要求和附件，便于学生根据专家意见优化项目材料和路演内容。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"专家意见\"。")
    add_info_block(doc, "前置条件", "管理员已创建专家评审包并分配专家，且专家已提交评审意见。")
    add_info_block(doc, "操作步骤", [
        "进入专家意见页面，查看意见列表。",
        "点击单条意见，查看专家反馈摘要和详细说明。",
        "若有附件，可下载查看专家提供的参考资料或标注文件。",
        "根据意见内容，在项目管理中修改对应材料，或在训练中心中调整路演和答辩策略。",
    ])
    add_info_block(doc, "注意事项", [
        "团队成员可查看专家意见的文字内容，但不可下载专家评审材料（代码权限控制：member 角色的 materials 字段被置空）。",
        "项目负责人可查看完整的专家意见和评审材料。",
        "专家意见按时间倒序排列，最新意见展示在最上方。",
    ])
    add_figure(doc, "5-5")

    # 5.7 资料归档
    add_heading_custom(doc, "5.7 资料归档", level=2)
    add_info_block(doc, "功能用途", "查看本组已通过审核并归档的资料及历史版本，支持按分类筛选和下载。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"资料归档\"。")
    add_info_block(doc, "前置条件", "材料已通过审核（approved 状态）。")
    add_info_block(doc, "操作步骤", [
        "进入资料归档页面，查看本组资料列表。",
        "在左侧分类栏中选择资料类型（如计划书、路演PPT、视频、证明附件等），筛选对应资料。",
        "点击资料名称，查看详情和历史版本。",
        "如需下载，点击下载按钮保存至本地。",
    ])
    add_info_block(doc, "注意事项", [
        "资料归档为只读视图，学生不能在归档页面直接上传新文件；如需更新材料，请在\"项目管理\"中重新提交。",
        "历史版本保留机制：同一材料多次上传后，系统保留历史版本，便于追溯变更过程。",
    ])
    add_placeholder_figure(doc, "图 5-6  学生端资料归档页面", "建议截取：左侧菜单\"资料归档\"入口、分类筛选栏、资料列表。")

    # 5.8 训练中心
    add_heading_custom(doc, "5.8 训练中心", level=2)
    add_info_block(doc, "功能用途", "提供模拟Q&A题库、随机抽查、路演计时和训练记录，帮助学生团队提升答辩应对能力和路演时间掌控能力。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"训练中心\"。")
    add_info_block(doc, "前置条件", "无特殊前置条件。")
    add_info_block(doc, "操作步骤", [
        "进入训练中心，选择训练模式：模拟问答或路演计时。",
        "模拟问答模式下，系统从题库中随机抽取问题，学生现场作答并记录表现。",
        "路演计时模式下，设置路演时长（如5分钟陈述+3分钟问答），系统倒计时并提示剩余时间。",
        "训练结束后，查看训练记录和历次表现对比。",
    ])
    add_info_block(doc, "注意事项", [
        "训练中心数据仅对本组成员可见，不会同步给专家或管理员。",
        "建议定期使用训练中心进行模拟演练，尤其在正式路演前一周保持高频训练。",
    ])
    add_placeholder_figure(doc, "图 5-7  学生端训练中心页面", "建议截取：左侧菜单\"训练中心\"入口、训练模式选择、计时器界面。")

    # 5.9 AI助手
    add_heading_custom(doc, "5.9 AI助手", level=2)
    add_info_block(doc, "功能用途", "基于大语言模型的智能助手，可咨询系统使用方法、赛事流程规则、材料撰写规范等问题。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"AI助手\"。")
    add_info_block(doc, "前置条件", "管理员已为本账号开通AI使用权限（默认按角色批量配置）。")
    add_info_block(doc, "操作步骤", [
        "进入AI助手页面，在输入框中描述需要咨询的问题。",
        "AI将基于系统知识库和赛事规则给出回答。",
        "可继续追问，形成多轮对话。",
    ])
    add_info_block(doc, "注意事项", [
        "AI助手的回答仅供参考，涉及重要决策（如材料提交、评分规则）请以管理员正式通知为准。",
        "AI使用配额由管理员统一配置，高频使用可能触发配额提醒。",
    ])
    add_placeholder_figure(doc, "图 5-8  学生端AI助手页面", "建议截取：左侧菜单\"AI助手\"入口、对话界面、示例问答。")

    # 5.10 个人信息维护
    add_heading_custom(doc, "5.10 个人信息维护", level=2)
    add_info_block(doc, "功能用途", "查看并维护当前登录账号的个人资料，包括头像、联系方式、密码修改等。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"个人信息\"。")
    add_info_block(doc, "前置条件", "无特殊前置条件。")
    add_info_block(doc, "操作步骤", [
        "查看个人基本信息：姓名、角色、项目组、学院、学号/工号等。",
        "点击头像区域上传新头像（支持常见图片格式）。",
        "修改邮箱或补充其他联系方式。",
        "在密码区域输入原密码和新密码，完成密码修改。",
        "点击\"保存\"使变更生效。",
    ])
    add_info_block(doc, "注意事项", [
        "部分敏感字段（如角色、项目组）仅管理员可修改，个人无法自行变更。",
        "修改密码后，建议使用新密码重新登录一次，确认生效。",
    ])
    add_placeholder_figure(doc, "图 5-9  学生端个人信息页面", "建议截取：左侧菜单\"个人信息\"入口、资料编辑表单。")

    doc.add_page_break()

    # ============================================================
    # 第六章 指导教师端操作办法
    # ============================================================
    add_heading_custom(doc, "第六章 指导教师端操作办法", level=1)
    add_paragraph_custom(doc,
        "指导教师负责查看所绑定项目组的日常进展、进行点评指导、审核项目材料，并协助学生落实专家意见。教师端与学生端在界面布局上基本一致，但在数据范围和部分管理权限上有所扩展。",
        first_line_indent=0.74)

    # 6.1 首页概览
    add_heading_custom(doc, "6.1 首页概览", level=2)
    add_info_block(doc, "功能用途", "查看负责项目组的整体进展、待处理任务、学生汇报情况和系统通知。")
    add_info_block(doc, "进入路径", "指导教师账号登录后自动进入；或点击左侧菜单\"首页概览\"。")
    add_info_block(doc, "前置条件", "账号已审核通过且已绑定项目组。")
    add_info_block(doc, "操作步骤", [
        "查看顶部通知和待办事项，及时处理学生提交或管理员分配的事项。",
        "查看负责项目组的进度节点卡片，关注材料提交截止时间和评审安排。",
        "通过快捷入口跳转至日程汇报、项目管理或任务工单模块。",
    ])
    add_info_block(doc, "注意事项", [
        "教师端首页统计范围为教师所绑定的全部项目组（若一名教师指导多个项目，可切换查看）。",
        "关注红色预警卡片，及时处理逾期任务和未提交日报的学生。",
    ])
    add_figure(doc, "6-1")

    # 6.2 日程汇报查看与点评
    add_heading_custom(doc, "6.2 日程汇报查看与点评", level=2)
    add_info_block(doc, "功能用途", "选择日期和项目组，查看学生日报提交情况，对日报进行点赞、点评或提出修改建议，并督促未提交学生补交。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"日程汇报\"。")
    add_info_block(doc, "前置条件", "学生已提交对应日期的日报。")
    add_info_block(doc, "操作步骤", [
        "在日程汇报页面，选择要查看的日期（默认为当天）。",
        "在项目组筛选器中，选择要查看的项目组（若指导多个项目组）。",
        "查看成员汇报卡片：已提交的成员显示汇报摘要；未提交的成员显示\"未提交\"状态。",
        "点击单条汇报，展开详情，查看今日完成工作、遇到的问题和下一步计划。",
        "发起点评：在汇报下方输入点评内容，选择点评类型（表扬/改进建议/普通评论），点击提交。",
        "督促补交：对未提交或连续异常的学生，点击\"发送提醒\"，系统将通知邮件发送给学生。",
    ])
    add_info_block(doc, "提交后流转", "点评即时保存并显示在学生的日程汇报页面中；督促提醒通过系统消息和邮件通知学生。")
    add_info_block(doc, "注意事项", [
        "教师可查看本组全部成员的汇报记录，但无法代学生填写或修改日报内容。",
        "点评建议应具体、可操作，便于学生明确改进方向。",
        "提醒功能发送的邮件内容由系统模板生成，教师可在发送前预览。",
    ])
    add_figure(doc, "6-2")

    # 6.3 项目管理（材料审核）
    add_heading_custom(doc, "6.3 项目管理（材料审核）", level=2)
    add_info_block(doc, "功能用途", "查看本组当前阶段材料提交情况，预览学生上传的PDF和视频，执行审核操作（通过或退回修改）。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"项目管理\"。")
    add_info_block(doc, "前置条件", "学生已提交项目材料。")
    add_info_block(doc, "操作步骤", [
        "进入项目管理页面，查看本组材料列表。",
        "点击材料名称或\"预览\"按钮，在线查看PDF计划书、路演PPT或播放项目视频。",
        "审核判断：材料符合要求时，点击\"通过\"；需要修改时，点击\"退回\"并填写具体修改意见。",
        "跟踪状态：审核通过的材料状态变为\"已通过\"，并自动进入资料归档；被退回的材料状态变为\"需修改\"，学生可在项目管理中心重新上传。",
    ])
    add_info_block(doc, "提交后流转", "pending → teacher_approved（教师通过）→ approved（终审通过） 或 revision（退回修改）。")
    add_info_block(doc, "注意事项", [
        "教师具备 `canTeacherReviewDocument` 权限，可对材料进行审核，但终审通过可能还需项目负责人确认（取决于系统配置）。",
        "退回材料时，意见描述应尽可能详细，指出具体问题和修改建议，减少反复退回。",
        "审核结果对学生端实时可见，建议及时审核以免影响学生后续工作安排。",
    ])
    add_figure(doc, "6-3")

    # 6.4 资料归档
    add_heading_custom(doc, "6.4 资料归档", level=2)
    add_info_block(doc, "功能用途", "查看本组过程材料、终审材料和已归档文件的历史版本，支持下载和预览。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"资料归档\"。")
    add_info_block(doc, "前置条件", "材料已通过审核。")
    add_info_block(doc, "操作步骤", [
        "进入资料归档页面，按分类筛选查看本组资料。",
        "点击资料名称查看详情，包括上传人、上传时间、审核记录和历史版本。",
        "如需本地备份，点击下载按钮保存文件。",
    ])
    add_info_block(doc, "注意事项", [
        "资料归档页面为只读，教师不能在此页面上传或修改文件。",
        "历史版本保留便于追溯，建议定期下载重要里程碑版本进行本地备份。",
    ])
    add_figure(doc, "6-4")

    # 6.5 任务工单
    add_heading_custom(doc, "6.5 任务工单", level=2)
    add_info_block(doc, "功能用途", "查看与本组相关的工单，分配学生处理或确认完成。教师具备完整的工单管理权限。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"任务中心\"。")
    add_info_block(doc, "前置条件", "无特殊前置条件。")
    add_info_block(doc, "操作步骤", [
        "查看工单看板，了解本组全部工单的状态分布。",
        "创建工单：点击\"新建工单\"，填写标题、描述、截止时间，分配处理人（学生），设置验收人。",
        "验收任务：学生在\"待验收\"列中提交的任务，教师查看完成说明和附件后，点击\"通过\"或\"退回\"。",
        "编辑/删除：对未归档的工单，教师可随时编辑内容或删除（已归档工单不可删除）。",
    ])
    add_info_block(doc, "权限说明", "教师可创建、编辑、删除工单，可验收学生提交的任务；与管理员相比，教师只能操作本组工单，不能查看或操作其他项目组的工单。")
    add_info_block(doc, "注意事项", [
        "删除工单为不可逆操作，删除前请确认工单内无重要附件或未完成的验收流程。",
        "建议优先使用系统分配功能明确责任人，避免口头传达导致的任务遗漏。",
    ])
    add_placeholder_figure(doc, "图 6-5  指导教师端任务工单页面", "建议截取：左侧菜单\"任务中心\"入口、工单看板、创建工单弹窗。")

    # 6.6 专家意见与整改指导
    add_heading_custom(doc, "6.6 专家意见与整改指导", level=2)
    add_info_block(doc, "功能用途", "查看本组收到的专家意见，逐条指导学生落实整改，跟踪整改完成情况。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"专家意见\"。")
    add_info_block(doc, "前置条件", "管理员已分配专家评审且专家已提交意见。")
    add_info_block(doc, "操作步骤", [
        "进入专家意见页面，查看专家反馈列表。",
        "阅读每条意见的具体内容，结合项目实际判断整改优先级。",
        "将整改要求分解为具体任务，通过\"任务中心\"分配给学生执行。",
        "跟踪整改进度，确认学生已在项目管理中更新对应材料或在训练中心中调整答辩策略。",
    ])
    add_info_block(doc, "注意事项", [
        "教师可查看专家上传的评审材料附件，便于更准确地理解专家意见。",
        "对争议性意见，教师可结合项目实际情况进行判断，必要时向管理员反馈。",
    ])
    add_placeholder_figure(doc, "图 6-6  指导教师端专家意见页面", "建议截取：左侧菜单\"专家意见\"入口、意见列表、附件下载区。")

    # 6.7 团队管理
    add_heading_custom(doc, "6.7 团队管理", level=2)
    add_info_block(doc, "功能用途", "查看本组成员分工、账号信息和角色配置。教师可查看本组成员，但不可删除账号，也不可审核注册申请。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"团队管理\"。")
    add_info_block(doc, "前置条件", "无特殊前置条件。")
    add_info_block(doc, "操作步骤", [
        "进入团队管理页面，查看本组成员列表。",
        "查看成员基本信息：姓名、角色、负责内容、今日日报摘要、任务完成进度等。",
        "点击成员卡片，查看详细资料和联系方式（若权限允许）。",
    ])
    add_info_block(doc, "权限边界", [
        "教师可查看本组成员信息，但\"删除账号\"入口对其不可见。",
        "教师不能审核注册申请（审核权限仅限系统管理员和校级管理员）。",
        "教师不能修改学生的角色或将其调整到其他项目组（此类操作需管理员执行）。",
    ])
    add_info_block(doc, "注意事项", [
        "团队管理页面的数据为实时同步，学生端的资料变更会即时反映在教师视图中。",
        "若发现成员信息有误，请联系管理员进行修改。",
    ])
    add_placeholder_figure(doc, "图 6-7  指导教师端团队管理页面", "建议截取：左侧菜单\"团队管理\"入口、成员列表卡片。")

    # 6.8 训练中心、AI助手、个人信息
    add_heading_custom(doc, "6.8 训练中心、AI助手与个人信息", level=2)
    add_paragraph_custom(doc,
        "上述三个模块的功能与学生端基本一致。指导教师可使用训练中心进行模拟答辩演练，使用AI助手咨询系统使用问题，并在个人信息中维护账号资料。差异点在于：教师端训练中心的数据同样仅对本组可见；AI助手的配额由管理员按教师角色统一配置。",
        first_line_indent=0.74)
    add_placeholder_figure(doc, "图 6-8  指导教师端训练中心页面", "建议截取：训练模式选择、计时器。")

    doc.add_page_break()

    # ============================================================
    # 第七章 管理员端操作办法
    # ============================================================
    add_heading_custom(doc, "第七章 管理员端操作办法", level=1)
    add_paragraph_custom(doc,
        "管理员包括系统管理员和校级管理员两类角色。两者在系统功能上高度一致，仅在账号删除权限上存在细微差异：系统管理员可删除任何非系统管理员账号；校级管理员可删除非校级管理员账号。本章统一以\"管理员\"身份进行描述，涉及差异处单独标注。",
        first_line_indent=0.74)

    # 7.1 首页概览
    add_heading_custom(doc, "7.1 首页概览", level=2)
    add_info_block(doc, "功能用途", "查看全校项目进展总览、待审核账号数量、未读消息、资料审批状态和任务工单统计。首页提供全校维度的关键指标和异常提醒。")
    add_info_block(doc, "进入路径", "管理员账号登录后自动进入；或点击左侧菜单\"首页概览\"。")
    add_info_block(doc, "前置条件", "无特殊前置条件。")
    add_info_block(doc, "操作步骤", [
        "查看全校概览卡片：包括总项目组数、总成员数、待审核账号数、今日日报提交率等。",
        "查看顶部操作区：发布公告、查看待办、查看消息、提交Bug反馈。",
        "关注异常提醒：未提交日报的项目组、逾期未完成的任务、材料审核积压情况等以红色或橙色标记提示。",
        "点击各统计卡片，可快速跳转至对应的详细管理页面。",
    ])
    add_info_block(doc, "注意事项", [
        "管理员首页为全校视角，数据范围覆盖所有项目组，不受单个项目组绑定限制。",
        "异常提醒按紧急程度排序，建议优先处理红色预警项。",
    ])
    add_figure(doc, "7-1")

    # 7.2 团队管理与账号审核
    add_heading_custom(doc, "7.2 团队管理与账号审核", level=2)
    add_info_block(doc, "功能用途", "统一管理全校账号、创建和维护项目组、审核自助注册申请、重置密码、删除账号。这是系统准入控制的核心模块。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"团队管理\"。")
    add_info_block(doc, "前置条件", "无特殊前置条件。")
    add_info_block(doc, "操作步骤", [
        "查看账号列表：页面展示全校全部账号，支持按角色、项目组、审核状态筛选。",
        "创建项目组：点击\"创建项目组\"，填写项目组名称、选择负责人和指导教师，保存后项目组生效。",
        "审核注册申请：在待审核区域查看 pending 状态的账号，点击\"审核\"。",
        "分配项目组：在审核弹窗中，必须为待审核账号选择一个项目组；不分配项目组则无法通过审核。",
        "通过/拒绝：确认信息无误且项目组已分配后，点击\"通过\"；若信息有误或不符合要求，点击\"拒绝\"并填写原因。",
        "账号维护：对已通过账号，可编辑角色、重置密码、调整所属项目组或删除账号。",
        "批量添加专家：点击\"批量添加专家\"，按模板格式导入专家名单，系统自动创建专家账号。",
    ])
    add_info_block(doc, "提交后流转", "审核通过后，账号状态变为\"已通过\"，用户可立即登录系统。被拒绝的账号可重新注册或联系管理员申诉。")
    add_info_block(doc, "权限边界", [
        "系统管理员可删除任何非系统管理员账号；校级管理员可删除非校级管理员账号。",
        "学生、指导教师、项目负责人、团队成员均不显示删除账号入口。",
        "只有系统管理员和校级管理员能审核注册申请和创建直属账号。",
        "只有系统管理员能将其他账号设为系统管理员；校级管理员不能创建或修改系统管理员账号。",
    ])
    add_info_block(doc, "注意事项", [
        "审核时必须分配项目组，这是系统强制规则，不可跳过。",
        "删除账号为不可逆操作，删除后该账号的所有日报、任务记录等数据仍保留在系统中，但账号无法登录。",
        "建议定期清理长期未登录的僵尸账号，保持团队管理列表清晰。",
        "批量添加专家时，请确保邮箱地址准确，系统将自动发送账号通知邮件。",
    ])
    add_figure(doc, "7-2")

    # 7.3 项目阶段与材料管理
    add_heading_custom(doc, "7.3 项目阶段与材料管理", level=2)
    add_info_block(doc, "功能用途", "创建或编辑项目阶段，配置材料要求、提交时间窗口、面向项目组和审核规则；按项目组查看材料提交、审核和归档状态。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"项目管理\"。")
    add_info_block(doc, "前置条件", "无特殊前置条件。")
    add_info_block(doc, "操作步骤", [
        "创建阶段：点击\"创建阶段\"，选择阶段类型。",
        "配置类型：选择\"网络评审\"（需上传材料）或\"项目路演\"（不要求上传材料，用于现场评审）。",
        "设置材料要求：网络评审阶段默认要求计划书PDF、PPT PDF和视频；可根据需要增删。",
        "设置面向项目组：选择该阶段对哪些项目组开放；不选择则默认对全部项目组开放。",
        "设置时间窗口：填写开始时间和截止时间，控制材料提交的开放时段。",
        "保存阶段：确认信息无误后保存，阶段即对学生和教师可见。",
        "查看提交状态：在阶段卡片下方查看各项目组的材料提交情况，对未按时提交的项目组进行提醒。",
    ])
    add_info_block(doc, "提交后流转", "阶段创建后立即生效，学生在\"项目管理\"中按阶段要求上传材料。管理员可随时编辑阶段信息或关闭阶段。")
    add_info_block(doc, "注意事项", [
        "\"项目路演\"阶段不要求上传材料，主要用于现场大屏评审的配置关联。",
        "截止时间到期后，学生将无法再提交材料（除非管理员延长截止时间）。",
        "阶段面向项目组的配置可随时调整，调整后新开放的项目组即可看到该阶段。",
        "对已有关联评审包的阶段，修改阶段类型可能导致评审包配置失效，请谨慎操作。",
    ])
    add_figure(doc, "7-3")

    # 7.4 专家评审组织
    add_heading_custom(doc, "7.4 专家评审组织", level=2)
    add_info_block(doc, "功能用途", "创建评审包、分配评审专家、设置评分规则（去最高分/最低分）和评审时间窗口，是组织专家评审工作的核心模块。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"专家评审\"。")
    add_info_block(doc, "前置条件", "已创建项目阶段；已创建专家账号（或在专家库中存在可用专家）。")
    add_info_block(doc, "操作步骤", [
        "创建评审包：点击\"创建评审包\"，选择关联的项目阶段。",
        "选择参评项目组：从项目组列表中勾选参与本轮评审的项目组。",
        "配置专家：从专家库中选择评审专家，一个评审包可分配多名专家。",
        "设置计分规则：配置去掉最高分和去掉最低分的数量（例如各去1个）。系统要求保留的有效评分专家数不少于2人。",
        "设置评审时间窗口：填写评审开始时间和截止时间。网评模式下，专家只能在时间窗口内提交评分。",
        "保存评审包：确认配置无误后保存，评审包状态变为\"已配置\"，专家端即可看到对应评审任务。",
        "上传评审材料（可选）：为评审包上传计划书、路演PPT和视频，供专家在线预览。",
    ])
    add_info_block(doc, "提交后流转", "评审包创建后，专家在专家评审页面看到分配的任务。专家完成评分后，管理员可在评审包详情中查看各专家评分和系统计算的最终得分。")
    add_info_block(doc, "注意事项", [
        "评审包一旦创建，不建议频繁修改专家分配，以免影响评审进度。",
        "计分规则中\"去掉最高分/最低分\"的数量之和必须小于专家总数减2，否则系统会提示规则无效。",
        "网评和路演的评审包可分别创建，但一个评审包只能关联一种阶段类型。",
        "评审材料上传后支持覆盖更新，但已提交评分的专家可能基于旧版材料打分，请谨慎更新。",
    ])
    add_figure(doc, "7-4")

    # 7.5 现场大屏控制（路演评审）
    add_heading_custom(doc, "7.5 现场大屏控制（路演评审）", level=2)
    add_info_block(doc, "功能用途", "在路演现场通过大屏控制台执行全流程评审管理，包括抽签排序、路演计时、答辩计时、评分控制和最终得分揭晓。")
    add_info_block(doc, "进入路径", "管理员在\"专家评审\"中打开评审包详情 → 点击\"现场大屏\"或\"生成大屏链接\"。")
    add_info_block(doc, "前置条件", "已创建路演类型的评审包；已分配评审专家；专家已就位。")
    add_info_block(doc, "操作步骤", [
        "生成大屏链接：在评审包中点击生成大屏链接，将链接投射至现场大屏幕。",
        "抽签排序（draw）：点击\"开始抽签\"，系统随机排列参评项目组顺序，大屏展示抽签结果。",
        "开始路演展示（presentation）：按抽签顺序，点击\"开始路演\"，系统启动倒计时（默认8分钟）。路演团队开始陈述。",
        "开始答辩提问（qa）：路演陈述结束后，点击\"开始答辩\"，系统启动答辩倒计时（默认7分钟）。评委提问，团队回答。",
        "开始评分（scoring）：答辩结束后，点击\"开始评分\"，系统通知所有专家进入评分状态，大屏显示匿名专家席位和评分进度。专家在各自设备上输入总分（0.00-100.00，保留两位小数）。",
        "确认并揭晓分数（reveal）：所有专家提交评分后，点击\"揭晓分数\"，系统根据预设规则（去掉最高/最低分）计算最终得分，大屏以动画形式展示最终得分。",
        "本轮结束（finished）：得分揭晓后，点击\"本轮结束\"，系统保存本轮结果，进入下一项目组或结束评审。",
    ])
    add_info_block(doc, "提交后流转", "每轮评审结果自动保存至系统后台，管理员可在评审包详情中查看历史得分记录和专家评分明细。")
    add_info_block(doc, "注意事项", [
        "大屏展示的专家席位为匿名形式（显示为\"专家1\"、\"专家2\"等），避免泄露专家真实身份。",
        "评分倒计时结束后，未提交的专家将无法再提交，系统自动标记为未提交状态。",
        "若某专家因故无法参与评分，管理员可在大屏控制台将其席位设为\"作废\"，该席位不计入最终得分计算。",
        "建议在正式评审前进行一次全流程彩排，确保大屏链接、倒计时和评分提交均正常。",
        "【不确定，需要人工确认】大屏各阶段按钮的具体文案和顺序可能因现场设备分辨率略有调整，请以实际系统界面为准。",
    ])
    add_placeholder_figure(doc, "图 7-7  现场大屏控制台", "建议截取：大屏控制端界面，包含阶段控制按钮（抽签、路演、答辩、评分、揭晓）、计时器、专家席位状态。")

    # 7.6 日程汇报统计
    add_heading_custom(doc, "7.6 日程汇报统计", level=2)
    add_info_block(doc, "功能用途", "按项目组查看日报提交率、趋势图、异常项目和教师点评情况，为管理者提供数据化 oversight 手段。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"日程汇报\"。")
    add_info_block(doc, "前置条件", "无特殊前置条件。")
    add_info_block(doc, "操作步骤", [
        "选择查看日期：在日期选择器中切换日期，查看全校各项目组在该日期的日报提交情况。",
        "项目组筛选：在项目组下拉框中选择特定项目组，聚焦查看。",
        "查看趋势图：页面展示近期日报提交率趋势，帮助管理者识别波动和异常。",
        "查看异常提醒：对连续未提交日报的项目组或个人，系统以红色标记提醒。",
        "发送督促：点击\"发送提醒\"，向指定成员或项目组发送邮件督促补交日报。",
    ])
    add_info_block(doc, "注意事项", [
        "管理员可查看全校全部项目组的日报数据，数据范围不受项目组绑定限制。",
        "趋势图的数据更新可能存在短暂延迟（通常不超过5分钟）。",
    ])
    add_figure(doc, "7-5")

    # 7.7 任务工单（全校任务台账）
    add_heading_custom(doc, "7.7 任务工单（全校任务台账）", level=2)
    add_info_block(doc, "功能用途", "以全校视角发布和管理任务工单，分配处理人，跟踪状态流转，实现跨项目组的任务协调。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"任务中心\"。")
    add_info_block(doc, "前置条件", "无特殊前置条件。")
    add_info_block(doc, "操作步骤", [
        "查看全校任务台账：页面以看板形式展示全校所有工单，支持按项目组、状态筛选。",
        "发布工单：点击\"新建工单\"，填写标题、描述、截止时间，选择处理人（全校成员可选）和验收人。",
        "分配与跟踪：工单创建后进入\"待分配/待接取\"列，处理人接取后进入\"处理中\"，完成后提交验收。",
        "验收与归档：管理员作为验收人审核任务完成质量，通过后状态变为\"已完成\"，可进一步归档。",
    ])
    add_info_block(doc, "注意事项", [
        "管理员发布的工单可分配给全校任何成员，不受项目组限制。",
        "工单删除为不可逆操作，删除前请确认无未完成的验收流程。",
        "建议对全校性任务（如系统培训、材料统一收集）使用工单管理，确保责任到人、状态可追踪。",
    ])
    add_figure(doc, "7-6")

    # 7.8 资料归档
    add_heading_custom(doc, "7.8 资料归档", level=2)
    add_info_block(doc, "功能用途", "按项目组查看全校已通过审核的资料和历史版本，支持下载、预览和审批记录追溯。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"资料归档\"。")
    add_info_block(doc, "前置条件", "无特殊前置条件。")
    add_info_block(doc, "操作步骤", [
        "进入资料归档页面，左侧展示全校项目组列表。",
        "点击项目组名称，展开该组的资料分类列表。",
        "按分类筛选（计划书、路演PPT、视频、证明附件等）查看具体资料。",
        "点击资料名称查看详情，包括历史版本、审核记录、上传人和上传时间。",
    ])
    add_info_block(doc, "注意事项", [
        "管理员在资料归档页面具备删除权限，可删除任何资料或历史版本。",
        "删除资料前请确认该资料未被关联的评审包引用，否则可能导致评审材料缺失。",
    ])
    add_placeholder_figure(doc, "图 7-8  管理员端资料归档页面", "建议截取：左侧项目组列表、分类筛选、资料详情弹窗。")

    # 7.9 时间进度管理
    add_heading_custom(doc, "7.9 时间进度管理", level=2)
    add_info_block(doc, "功能用途", "创建、编辑和删除赛事关键时间节点，向全校师生统一发布时间安排。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"时间进度\"。")
    add_info_block(doc, "前置条件", "无特殊前置条件。")
    add_info_block(doc, "操作步骤", [
        "查看现有事件列表，按时间远近排序。",
        "创建事件：点击\"新建事件\"，填写事件名称、日期时间、地点和说明，保存后事件对所有用户可见。",
        "编辑/删除：点击事件卡片上的编辑或删除按钮，修改事件信息或移除事件。",
    ])
    add_info_block(doc, "权限说明", "系统管理员、校级管理员和指导教师均可编辑时间进度；学生端为只读。")
    add_info_block(doc, "注意事项", [
        "事件时间修改后，系统不会自动通知已关注该事件的用户，建议通过公告或工单另行告知。",
        "关键节点（如材料截止、路演日期）建议提前在系统中录入，便于首页倒计时展示。",
    ])
    add_placeholder_figure(doc, "图 7-9  管理员端时间进度页面", "建议截取：事件列表、新建事件弹窗。")

    # 7.10 公告发布与系统反馈
    add_heading_custom(doc, "7.10 公告发布与系统反馈", level=2)
    add_info_block(doc, "功能用途", "发布全校公告、查看和处理用户提交的Bug反馈。")
    add_info_block(doc, "进入路径", "公告发布：首页顶部操作区 → \"发布公告\"；Bug反馈：顶部帮助/反馈入口。")
    add_info_block(doc, "前置条件", "仅系统管理员和校级管理员可发布公告。")
    add_info_block(doc, "操作步骤", [
        "发布公告：填写公告标题和内容，选择发布范围（全校），点击发布。公告将显示在所有用户的首页通知区。",
        "查看反馈：点击顶部帮助图标 → \"查看反馈\"，查看用户提交的Bug列表。",
        "处理反馈：阅读问题描述，定位原因后修复或在系统中回复用户。",
    ])
    add_info_block(doc, "注意事项", [
        "公告发布后即时生效，建议发布前仔细核对内容，避免信息错误引起误解。",
        "Bug反馈由系统管理员在消息中收到，建议建立内部处理流程，确保用户问题得到及时响应。",
    ])
    add_placeholder_figure(doc, "图 7-10  公告发布页面", "建议截取：发布公告弹窗、标题和内容编辑器。")

    # 7.11 AI权限配置
    add_heading_custom(doc, "7.11 AI权限配置", level=2)
    add_info_block(doc, "功能用途", "按成员或按角色批量配置AI助手的使用权限和配额，控制AI资源消耗。")
    add_info_block(doc, "进入路径", "左侧菜单 → \"团队管理\" → 点击成员列表中的AI权限配置入口。")
    add_info_block(doc, "前置条件", "无特殊前置条件。")
    add_info_block(doc, "操作步骤", [
        "在团队管理页面，筛选需要配置AI权限的成员。",
        "勾选成员后，点击\"批量配置AI权限\"，设置是否允许使用AI助手及每日/每月配额上限。",
        "保存配置后，被配置成员即可在AI助手页面使用对应配额。",
    ])
    add_info_block(doc, "权限说明", "仅系统管理员和校级管理员可见AI权限配置入口。")
    add_info_block(doc, "注意事项", [
        "AI配额为系统级资源，建议根据角色重要性合理分配，避免个别用户过度消耗。",
        "配额修改即时生效，已超限的用户在下次对话时将收到配额不足提示。",
    ])
    add_placeholder_figure(doc, "图 7-11  AI权限配置页面", "建议截取：批量配置弹窗、配额滑块。")

    doc.add_page_break()

    # ============================================================
    # 第八章 专家评审端操作办法
    # ============================================================
    add_heading_custom(doc, "第八章 专家评审端操作办法", level=1)
    add_paragraph_custom(doc,
        "评审专家为独立角色，系统界面与普通用户显著不同。专家登录后仅可见\"专家评审\"和\"个人信息\"两个模块，确保专家专注于评审任务，避免无关信息干扰。",
        first_line_indent=0.74)

    # 8.1 专家登录与首页
    add_heading_custom(doc, "8.1 专家登录与首页", level=2)
    add_info_block(doc, "功能用途", "评审专家通过系统地址登录，进入独立的专家评审入口。首页展示当前专家姓名、所属评审包列表和评审状态总览。")
    add_info_block(doc, "进入路径", "浏览器访问系统地址 → 使用管理员分配的专家账号和密码登录。")
    add_info_block(doc, "前置条件", "管理员已创建专家账号，并已将该专家分配至至少一个评审包。")
    add_info_block(doc, "操作步骤", [
        "在登录页输入专家账号和密码。",
        "输入图形验证码，点击登录。",
        "登录后进入专家评审首页，查看\"我的评审任务\"列表。",
    ])
    add_info_block(doc, "注意事项", [
        "专家账号由管理员直接创建，不支持自助注册。",
        "若登录后提示\"暂无评审任务\"，请联系管理员确认是否已分配评审包。",
    ])
    add_placeholder_figure(doc, "图 8-1  专家端登录与首页", "建议截取：专家评审入口界面、评审任务列表。")

    # 8.2 网络评审操作
    add_heading_custom(doc, "8.2 网络评审操作", level=2)
    add_info_block(doc, "功能用途", "在线查看项目材料，根据所属项目类别自行参考打分，保留两位小数。")
    add_info_block(doc, "进入路径", "专家评审首页 → 点击待评审的评审包卡片 → 进入评分详情页。")
    add_info_block(doc, "前置条件", "评审包状态为\"已配置\"，当前时间在评审开始时间和截止时间之间（评审窗口期内）。")
    add_info_block(doc, "操作步骤", [
        "查看评审包信息：确认项目名称、评审轮次、概述和截止时间。",
        "在线预览材料：点击\"计划书\"、\"路演材料\"或\"视频\"，在系统内直接预览PDF文件或播放视频。",
        "输入评分：在评分输入框中输入总分，范围为 0.00 至 100.00，最多保留两位小数。",
        "填写总评意见：在评论框中填写对项目的综合评价、亮点和不足（可选）。",
        "提交评分：确认分数和意见无误后，点击\"提交\"。",
    ])
    add_info_block(doc, "提交后流转", "提交成功后，评审状态变为\"已提交\"。在评审窗口期关闭前，专家可修改评分；窗口期关闭或管理员锁定后，评分不可再修改。")
    add_info_block(doc, "注意事项", [
        "网评不区分具体打分维度，专家根据所属项目类别自行参考打分，保留两位小数。",
        "分数需为 0.00-100.00 之间的数值，系统不接受负数或超过100分的输入。",
        "评分必须在管理员设定的评审窗口期内完成，逾期将无法提交。",
        "后台模型兼容旧版四维度字段，但当前界面已统一为总分输入，专家无需关注维度拆分。",
    ])
    add_placeholder_figure(doc, "图 8-2  专家端网络评审评分页", "建议截取：评审包信息、材料预览区、评分输入框（0.00-100.00）、总评意见框、提交按钮。")

    # 8.3 路演评审操作
    add_heading_custom(doc, "8.3 路演评审操作", level=2)
    add_info_block(doc, "功能用途", "在现场路演环节通过大屏或独立页面实时打分，评分方式与网评一致，保留两位小数。")
    add_info_block(doc, "进入路径", "现场大屏评分开始后，专家在自己的设备上打开评分链接，或从专家评审首页进入对应路演评审包。")
    add_info_block(doc, "前置条件", "管理员已在大屏控制台点击\"开始评分\"，且当前评审包为路演类型。")
    add_info_block(doc, "操作步骤", [
        "等待管理员在大屏控制台启动评分阶段，专家页面自动解锁或刷新显示评分入口。",
        "进入评分页面，输入总分（0.00-100.00，保留两位小数）。",
        "确认分数后点击\"提交\"。",
        "等待大屏揭晓分数，系统根据评分规则自动计算并展示最终得分。",
    ])
    add_info_block(doc, "提交后流转", "提交成功后，大屏控制台实时显示该专家席位状态为\"已提交\"。所有专家提交后，管理员可触发揭晓分数。")
    add_info_block(doc, "注意事项", [
        "路演评分与网络评审使用同一套评分输入方式（总分 0.00-100.00，保留两位小数）。",
        "倒计时结束后未提交的评分将无法补交，请务必在限时内完成。",
        "专家在路演评分环节无需上传任何材料，仅负责输入分数。",
    ])
    add_placeholder_figure(doc, "图 8-3  专家端路演评分页", "建议截取：评分输入界面、倒计时提示、提交按钮。")

    # 8.4 查看已完成的评审
    add_heading_custom(doc, "8.4 查看已完成的评审", level=2)
    add_info_block(doc, "功能用途", "查看历史评分记录、评审状态和最终得分，便于专家回顾和核对。")
    add_info_block(doc, "进入路径", "专家评审首页 → 在评审包列表中筛选\"已提交\"或\"已结束\"状态的评审包。")
    add_info_block(doc, "前置条件", "专家已完成评分且评审包已结束。")
    add_info_block(doc, "操作步骤", [
        "在评审包列表中，点击已完成的评审包。",
        "查看自己提交的分数和总评意见。",
        "若评审包已揭晓最终得分，可查看系统计算后的最终得分和去分规则说明。",
    ])
    add_info_block(doc, "注意事项", [
        "已锁定的评审包不支持再次修改评分。",
        "专家无法查看其他专家的评分明细，仅能看到自己的评分和系统最终得分。",
    ])

    doc.add_page_break()

    # ============================================================
    # 第九章 常见问题与故障处理
    # ============================================================
    add_heading_custom(doc, "第九章 常见问题与故障处理", level=1)

    add_heading_custom(doc, "9.1 登录类问题", level=2)
    qa = [
        ("验证码看不清怎么办？", "点击验证码图片即可刷新，生成新的图形验证码后重新输入。"),
        ("忘记密码如何找回？", "在登录页点击\"忘记密码\"，输入注册邮箱，系统将发送重置链接至该邮箱。点击链接按提示设置新密码即可。"),
        ("提示\"登录状态已失效\"怎么办？", "系统会话有一定有效期，长时间未操作或网络异常会导致会话失效。请重新登录，建议及时保存重要数据。"),
    ]
    for q, a in qa:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"Q：{q}")
        set_run_font(run, "黑体", 10.5, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"A：{a}")
        set_run_font(run, "宋体", 10.5)

    add_heading_custom(doc, "9.2 注册与审核类问题", level=2)
    qa = [
        ("注册后为什么不能登录？", "注册后账号处于\"待审核\"状态，需要系统管理员或校级管理员审核通过并分配项目组后才能登录。这是系统的强制准入规则。"),
        ("审核等待时间太长怎么办？", "建议主动联系所在学院的管理员，提醒其尽快处理待审核账号。管理员在\"团队管理\"首页可直接看到待审核数量提醒。"),
        ("注册时提示\"账号名或邮箱已存在\"怎么办？", "说明该账号名或邮箱已被注册。请更换账号名或邮箱后重新提交。若确认邮箱属于自己且被他人冒用，请联系管理员处理。"),
    ]
    for q, a in qa:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"Q：{q}")
        set_run_font(run, "黑体", 10.5, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"A：{a}")
        set_run_font(run, "宋体", 10.5)

    add_heading_custom(doc, "9.3 权限与菜单类问题", level=2)
    qa = [
        ("为什么看不到某个菜单？", "不同角色的可见菜单不同。请确认当前账号角色是否具备该权限，以及账号是否已完成审核并绑定项目组。未绑定项目组的账号部分功能受限。"),
        ("为什么某个按钮是灰色的？", "灰色按钮代表当前角色不具备该操作权限，或当前数据状态不满足操作条件（如非开放时间、非本人任务等）。"),
        ("项目负责人和团队成员的权限差异是什么？", "项目负责人可创建和编辑工单、审核材料、查看本组全部日报；团队成员只能接取工单、提交日报、上传材料，不能创建工单或审核他人材料。"),
    ]
    for q, a in qa:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"Q：{q}")
        set_run_font(run, "黑体", 10.5, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"A：{a}")
        set_run_font(run, "宋体", 10.5)

    add_heading_custom(doc, "9.4 材料与评审类问题", level=2)
    qa = [
        ("材料上传提示格式不符合要求怎么办？", "计划书和路演PPT必须上传PDF格式，请使用Word或PowerPoint的\"另存为PDF\"功能导出后再上传。视频支持mp4、mov、avi格式。"),
        ("上传时提示文件太大怎么办？", "网络评审阶段视频大小限制为20MB。若文件过大，请使用视频压缩工具处理后重新上传，或联系管理员调整限制（如政策允许）。"),
        ("材料审核未通过如何重新提交？", "在项目管理页面，找到被退回的材料，查看退回意见，修改后点击\"重新上传\"，选择修改后的文件再次提交。"),
        ("专家意见中的材料为什么下载不了？", "团队成员角色受权限限制，无法下载专家评审材料，仅可查看文字意见。项目负责人和指导教师可下载完整材料。"),
    ]
    for q, a in qa:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"Q：{q}")
        set_run_font(run, "黑体", 10.5, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"A：{a}")
        set_run_font(run, "宋体", 10.5)

    add_heading_custom(doc, "9.5 大屏与专家类问题", level=2)
    qa = [
        ("大屏链接打不开或显示异常怎么办？", "请检查现场网络连接是否正常，确认链接地址完整无误。若仍无法打开，请管理员在专家评审页面重新生成大屏链接。"),
        ("专家无法进入评分页面怎么办？", "请管理员确认已在大屏控制台点击\"开始评分\"，且专家设备网络正常。路演评分必须在管理员启动评分阶段后才会解锁。"),
        ("最终得分未生成或显示异常怎么办？", "请检查所有有效专家席位是否均已提交评分，以及是否有席位被误设为\"作废\"。若仍有异常，请联系系统管理员后台核查。"),
        ("专家评分提交后还能修改吗？", "在评审窗口期内且未被管理员锁定前，专家可修改评分。窗口期关闭或管理员手动锁定后，评分不可修改。"),
    ]
    for q, a in qa:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"Q：{q}")
        set_run_font(run, "黑体", 10.5, bold=True)
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Cm(0.8)
        run = p.add_run(f"A：{a}")
        set_run_font(run, "宋体", 10.5)

    add_heading_custom(doc, "9.6 系统问题反馈", level=2)
    add_paragraph_custom(doc,
        "如遇到本章节未覆盖的问题，可通过系统顶部帮助/反馈入口提交Bug反馈。填写问题标题和详细描述后提交，系统管理员将在消息中心收到反馈并进行处理。建议反馈时附上操作步骤、截图和出现时间，以便管理员快速定位和修复。",
        first_line_indent=0.74)

    doc.add_page_break()

    # ============================================================
    # 第十章 对外发布脱敏说明
    # ============================================================
    add_heading_custom(doc, "第十章 对外发布脱敏说明", level=1)

    add_heading_custom(doc, "10.1 脱敏范围与规则", level=2)
    add_paragraph_custom(doc,
        "本手册配套截图和示例数据已按以下规则完成脱敏处理，确保对外发布时不泄露真实个人信息和项目敏感内容：",
        first_line_indent=0.74)
    de_rules = [
        ("真实姓名", "统一替换为化名，如\"张三\"→\"学生A\"、\"李四\"→\"教师甲\"。"),
        ("账号名", "替换为示例账号或掩码，如\"zhangsan2023\"→\"student_2023****\"。"),
        ("学号", "保留前4位和后2位，中间用星号掩码，如\"2023010101\"→\"2023****01\"。"),
        ("工号", "同学号处理规则，保留前2位和后2位，中间掩码。"),
        ("邮箱地址", "替换为示例邮箱，如\"zhangsan@njt.edu.cn\"→\"student@example.edu.cn\"。"),
        ("手机号码", "保留前3位和后4位，中间用星号掩码，如\"13812345678\"→\"138****5678\"。"),
        ("项目名称", "替换为示例项目名称，如\"基于深度学习的轨道交通巡检系统\"→\"智慧交通创新项目\"。"),
        ("专家姓名", "匿名化处理，在系统界面和截图中显示为\"专家1\"、\"专家2\"等。"),
        ("学校名称", "本系统为南京铁道职业技术学院定制开发，对外发布时可根据实际情况决定是否保留。"),
    ]
    for item, rule in de_rules:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.left_indent = Cm(0.5)
        run = p.add_run(f"{item}：")
        set_run_font(run, "黑体", 10.5, bold=True)
        run = p.add_run(rule)
        set_run_font(run, "宋体", 10.5)

    add_heading_custom(doc, "10.2 截图脱敏检查清单", level=2)
    add_paragraph_custom(doc, "在将本手册截图替换为正式环境截图时，请对照以下清单逐项检查，确保所有敏感信息已脱敏：", first_line_indent=0.74)
    checks = [
        "检查截图中是否出现真实姓名（顶部用户栏、成员列表、专家列表等）。",
        "检查截图中是否出现真实学号、工号、手机号或邮箱。",
        "检查截图中是否出现真实项目名称（首页卡片、项目管理页、评审包页等）。",
        "检查截图中是否出现专家真实姓名（专家评审页、大屏控制台等）。",
        "检查浏览器地址栏是否暴露真实域名或IP地址。",
        "检查截图中的通知消息、邮件预览是否包含敏感内容。",
    ]
    for c in checks:
        add_bullet_paragraph(doc, c)

    add_heading_custom(doc, "10.3 文档发布审批流程", level=2)
    add_numbered_paragraph(doc, "初稿编制：由系统运维团队或项目管理团队完成初稿编写和截图嵌入。")
    add_numbered_paragraph(doc, "脱敏复核：由信息安全专员对照10.2节检查清单进行逐项复核，确认无敏感信息泄露。")
    add_numbered_paragraph(doc, "业务确认：由校级管理员或系统管理员确认操作流程和权限边界描述与实际系统一致。")
    add_numbered_paragraph(doc, "领导审批：提交学校分管领导审批，确认文档可用于对外汇报和正式发布。")
    add_numbered_paragraph(doc, "正式发布：审批通过后，以学校名义正式发布，版本号同步更新。")

    # ============================================================
    # 培训建议（附录）
    # ============================================================
    doc.add_page_break()
    add_heading_custom(doc, "附录：培训建议", level=1)
    add_paragraph_custom(doc,
        "正式培训时建议按以下顺序进行演示，确保不同角色的用户都能快速理解系统使用逻辑：",
        first_line_indent=0.74)
    add_numbered_paragraph(doc, "学生端演示（15分钟）：每日汇报填写与提交、项目管理材料上传、任务工单接取与验收。")
    add_numbered_paragraph(doc, "指导教师端演示（10分钟）：日程汇报点评、项目材料审核、任务工单创建与验收。")
    add_numbered_paragraph(doc, "管理员端演示（20分钟）：账号审核与项目组分配、项目阶段创建、专家评审包配置、现场大屏控制全流程。")
    add_numbered_paragraph(doc, "专家端演示（10分钟）：评审任务查看、材料预览、总分输入与提交。")
    add_numbered_paragraph(doc, "互动答疑（15分钟）：开放提问，重点解答权限边界、材料格式和评审规则问题。")

    add_paragraph_custom(doc,
        "建议培训前准备好测试账号（学生、教师、管理员、专家各至少一个），并确保测试环境中的项目阶段、评审包和任务工单已预置数据，以便现场演示时有真实案例可供操作。",
        first_line_indent=0.74)

    # 保存
    doc.save(OUTPUT_PATH)
    print(f"文档已生成：{OUTPUT_PATH}")


if __name__ == "__main__":
    build_document()
